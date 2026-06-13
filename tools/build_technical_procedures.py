from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "operaciones"
LOGO = ROOT / "assets" / "aci-logo-black.jpeg"

INK = "0B0F12"
WHITE = "FFFFFF"
AMBER = "C88C3A"
SEA = "1F5867"
STEEL = "5D788A"
PAPER = "F6F7F4"
LINE = "D7DDE1"
LIGHT = "E8EEF0"
CAUTION = "FFF4E5"


def set_font(run, size=10.3, color="25313A", bold=False, italic=False, name="Arial"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    node = tbl_pr.first_child_found_in("w:tblBorders")
    if node is None:
        node = OxmlElement("w:tblBorders")
        tbl_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = node.find(qn(f"w:{edge}"))
        if item is None:
            item = OxmlElement(f"w:{edge}")
            node.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), size)
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def base_doc(code, short_name):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.3)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, SEA, 14, 7),
        ("Heading 2", 13, SEA, 10, 5),
        ("Heading 3", 11.5, STEEL, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    set_font(header.add_run(f"ACI LATAM | {code} | {short_name}"), size=8, color=STEEL, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Version 1.0 | Para aprobacion e implementacion piloto"), size=8, color=STEEL)
    return doc


def paragraph(doc, text, bold=False, italic=False, color="25313A"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.1
    set_font(p.add_run(text), bold=bold, italic=italic, color=color)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    set_font(p.add_run(text), size=10.1)
    return p


def numbered(doc, title, body):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(f"{title}: "), bold=True)
    set_font(p.add_run(body))


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    borders(t)
    repeat_header(t.rows[0])
    for i, header in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Inches(widths[i])
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_margins(c)
        shade(c, INK)
        set_font(c.paragraphs[0].add_run(header), size=8.6, color=WHITE, bold=True)
    for row_index, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]
            c.width = Inches(widths[i])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_margins(c)
            if row_index % 2:
                shade(c, PAPER)
            set_font(c.paragraphs[0].add_run(str(value)), size=8.7, bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def callout(doc, label, text, fill=LIGHT):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.72)
    borders(t)
    c = t.cell(0, 0)
    shade(c, fill)
    cell_margins(c, top=130, start=160, bottom=130, end=160)
    p = c.paragraphs[0]
    set_font(p.add_run(f"{label}: "), color=SEA, bold=True)
    set_font(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def cover(doc, code, title, subtitle):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.72)
    borders(t, INK, "0")
    c = t.cell(0, 0)
    shade(c, INK)
    cell_margins(c, top=330, start=280, bottom=330, end=280)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(3.0))
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(22)
    set_font(p2.add_run(title.upper()), size=22, color=WHITE, bold=True)
    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p3.add_run(subtitle), size=11.5, color="DCE3E7", bold=True)
    p4 = c.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(16)
    set_font(p4.add_run(code), size=11, color=AMBER, bold=True)
    doc.add_paragraph()
    table(doc, ["Version", "Estado", "Fecha", "Aprobacion"], [["1.0", "Para aprobacion e implementacion piloto", "12-06-2026", "Pendiente"]], [0.8, 3.45, 1.1, 1.37])
    paragraph(doc, "Debe aplicarse junto con ACI-MO-001, la orden de servicio y las instrucciones contractuales vigentes. No reemplaza normas, metodos ni requisitos legales aplicables.", italic=True, color=STEEL)
    doc.add_page_break()


def common_start(doc, code, objective, scope):
    doc.add_heading("Control del documento", level=1)
    table(doc, ["Version", "Fecha", "Cambio", "Responsable"], [["1.0", "12-06-2026", "Emision inicial para aprobacion y prueba piloto.", "Operaciones / Calidad"]], [0.8, 1.1, 3.65, 1.17])
    doc.add_heading("1. Objeto", level=1)
    paragraph(doc, objective)
    doc.add_heading("2. Alcance", level=1)
    paragraph(doc, scope)
    doc.add_heading("3. Documentos relacionados", level=1)
    for item in ["ACI-MO-001 Manual Operacional de Loss Control y Expediting.", "ACI-F-001 Orden de Servicio.", "ACI-F-002 Checklist Operacional.", "ACI-F-003 Bitacora de Eventos.", "ACI-F-004 Registro Fotografico.", "ACI-F-005 Reporte Preliminar.", "ACI-F-006 Informe Final."]:
        bullet(doc, item)
    callout(doc, "Jerarquia", "La orden de servicio y la instruccion escrita del cliente definen el alcance. Si existe conflicto con seguridad, ley o integridad tecnica, se debe detener la actividad propia y escalar.", CAUTION)


def common_end(doc, records, pilot):
    doc.add_heading("Registros asociados", level=1)
    for item in records:
        bullet(doc, item)
    doc.add_heading("Implementacion piloto", level=1)
    paragraph(doc, pilot)
    for item in ["Capacitar al personal involucrado.", "Ejecutar un caso simulado con datos controlados.", "Registrar hallazgos y ajustes requeridos.", "Aprobar la version vigente antes del uso comercial."]:
        bullet(doc, item)


def quantity_doc():
    code = "ACI-PT-001"
    doc = base_doc(code, "Cantidad y Reconciliacion")
    cover(doc, code, "Procedimiento Tecnico", "Control de Cantidad y Reconciliacion")
    common_start(doc, code, "Definir el metodo para observar, registrar, comprobar y reconciliar cantidades en operaciones petroleras, manteniendo trazabilidad de datos, unidades, fuentes, calculos y limitaciones.", "Aplica a mediciones de nave, terminal, planta, tanques, unidades de transporte y documentos relacionados, cuando formen parte expresa del servicio contratado.")
    doc.add_heading("4. Principios de medicion y calculo", level=1)
    table(doc, ["Principio", "Aplicacion"], [
        ["Fuente identificada", "Cada cifra se marca como observada por ACI, calculada por ACI o suministrada por terceros."],
        ["Metodo definido", "Se utiliza el metodo, norma, tabla y base contractual especificados para la operacion."],
        ["Unidades consistentes", "No se comparan cifras hasta confirmar unidad, base de temperatura, densidad y condicion estandar."],
        ["Reproducibilidad", "Un tercero competente debe poder reconstruir el calculo con los datos y factores registrados."],
        ["Doble revision", "Toda diferencia relevante o cifra usada en reclamo debe ser revisada independientemente."],
    ], [1.45, 5.27])
    doc.add_heading("5. Preparacion", level=1)
    for item in [
        "Confirmar producto, cantidad estimada, unidades, base de referencia y tolerancia contractual.",
        "Obtener tablas de calibracion, certificados, factores, hojas de medicion y documentos disponibles.",
        "Identificar tanques, compartimientos, lineas, medidores y puntos de transferencia incluidos.",
        "Confirmar equipos que ACI observara o utilizara y su evidencia de calibracion/verificacion.",
        "Preparar hoja de calculo o formato controlado sin formulas ocultas o factores no documentados.",
    ]:
        bullet(doc, item)
    doc.add_heading("6. Control de mediciones", level=1)
    table(doc, ["Dato", "Control minimo"], [
        ["Identificacion", "Tanque/compartimiento, producto, fecha, hora, zona horaria y responsable."],
        ["Nivel o vacio", "Lectura, unidad, punto de referencia, repeticion y discrepancia entre lecturas."],
        ["Agua libre", "Metodo, punto, resultado, unidad y limitacion observada."],
        ["Temperatura", "Puntos, lecturas, promedio, instrumento y estabilidad."],
        ["Densidad", "Valor, temperatura de referencia, fuente y certificado relacionado."],
        ["Tabla/factor", "Identificacion, version, rango aplicable y fuente."],
        ["Sellos", "Numero, condicion, apertura/cierre y responsable cuando aplique."],
    ], [1.3, 5.42])
    doc.add_heading("7. Medicion en tanques de tierra", level=1)
    for item in [
        "Confirmar que el tanque y la tabla corresponden entre si.",
        "Registrar movimientos simultaneos, lineas abiertas, recirculacion, drenajes o transferencias que puedan afectar el balance.",
        "Observar condiciones de reposo y estabilidad indicadas por el metodo aplicable.",
        "Documentar lecturas iniciales y finales, correcciones y volumen neto calculado.",
        "Registrar limitaciones de acceso, espuma, turbulencia, techo flotante, sedimentos u otras condiciones relevantes.",
    ]:
        bullet(doc, item)
    doc.add_heading("8. Medicion en nave o unidad", level=1)
    for item in [
        "Identificar compartimientos incluidos y condición de los no incluidos.",
        "Registrar lecturas, temperaturas, agua, remanentes, slops y cantidades no bombeables cuando corresponda.",
        "Verificar que correcciones de trim/list u otras se apliquen desde fuente autorizada.",
        "Distinguir cifras de libro, cifras calculadas por terceros y cifras comprobadas por ACI.",
        "No presentar cifras de nave como independientes si ACI no observo las mediciones.",
    ]:
        bullet(doc, item)
    doc.add_heading("9. Reconciliacion", level=1)
    paragraph(doc, "La base de reconciliacion debe definirse antes de calcular diferencias. Las siguientes expresiones son estructuras generales y deben adaptarse a la base contractual:")
    table(doc, ["Concepto", "Estructura general"], [
        ["Cantidad entregada por origen", "Inventario inicial - inventario final +/- movimientos documentados."],
        ["Cantidad recibida por destino", "Inventario final - inventario inicial +/- movimientos documentados."],
        ["Diferencia", "Cantidad de destino - cantidad de origen, con signo y unidad claramente definidos."],
        ["Diferencia porcentual", "Diferencia / base de referencia acordada x 100."],
        ["Balance", "Entradas - salidas - variacion de inventario, incorporando movimientos autorizados."],
    ], [2.15, 4.57])
    callout(doc, "Prohibicion", "No mezclar volúmenes observados y estandar, masas y volumenes, o bases de temperatura distintas. No elegir una base solo porque reduce o aumenta la diferencia.", CAUTION)
    doc.add_heading("10. Investigacion de diferencias", level=1)
    numbered(doc, "Confirmar", "Revisar transcripcion, unidades, signos, decimales, formulas y documentos.")
    numbered(doc, "Aislar", "Comparar por tanque, compartimiento, periodo y punto de transferencia.")
    numbered(doc, "Evaluar", "Revisar agua, temperatura, densidad, remanentes, line fill, movimientos simultaneos, medidores y documentos pendientes.")
    numbered(doc, "Documentar", "Registrar hipotesis como hipotesis, indicando evidencia a favor, en contra y datos faltantes.")
    numbered(doc, "Escalar", "Comunicar diferencias superiores al criterio contractual o sin explicacion razonable.")
    doc.add_heading("11. Revision tecnica", level=1)
    for item in ["Cifras fuente verificadas contra documentos.", "Formulas visibles y reproducibles.", "Unidades y base de referencia consistentes.", "Factores y tablas identificados.", "Diferencias recalculadas por segundo revisor.", "Limitaciones y datos de terceros claramente indicados."]:
        bullet(doc, item)
    common_end(doc, ["ACI-F-002 Checklist Operacional.", "ACI-F-003 Bitacora.", "Hojas de medicion y calculo del servicio.", "ACI-F-005 Reporte Preliminar.", "ACI-F-006 Informe Final."], "Simular una reconciliacion con mediciones iniciales/finales, movimientos simultaneos, un dato incorrecto y una diferencia superior a tolerancia para verificar deteccion, escalamiento y doble revision.")
    return code, "ACI-PT-001-Control-Cantidad-Reconciliacion.docx", doc


def quality_doc():
    code = "ACI-PT-002"
    doc = base_doc(code, "Muestreo y Calidad")
    cover(doc, code, "Procedimiento Tecnico", "Muestreo, Calidad y Cadena de Custodia")
    common_start(doc, code, "Establecer controles para observar el muestreo, identificar muestras, mantener cadena de custodia y revisar documentos de calidad sin exceder la competencia ni el alcance contratado.", "Aplica a combustibles, crudos y derivados cuando ACI supervise o documente muestreo, sellado, traslado, recepcion de laboratorio o resultados de calidad.")
    doc.add_heading("4. Limites del servicio", level=1)
    callout(doc, "Limite tecnico", "ACI no declara conformidad de laboratorio ni sustituye al laboratorio acreditado. ACI documenta el proceso observado y compara resultados con la especificacion recibida cuando ello forme parte del alcance.", CAUTION)
    doc.add_heading("5. Plan de muestreo", level=1)
    for item in ["Producto, lote, tanque, nave o punto de transferencia.", "Objetivo del muestreo y ensayos requeridos.", "Metodo/norma/instruccion aplicable y responsable de ejecutarlo.", "Tipo, numero, volumen y distribucion de muestras.", "Recipientes, preservacion, sellos, etiquetas, transporte y laboratorio.", "Muestras de retencion, contramuestras y plazo de custodia."]:
        bullet(doc, item)
    doc.add_heading("6. Verificacion previa", level=1)
    table(doc, ["Control", "Criterio"], [
        ["Recipiente", "Limpio, compatible, integro, capacidad adecuada y cierre correcto."],
        ["Equipo", "Identificado, limpio, adecuado y en condicion de uso."],
        ["Punto", "Corresponde al lote/tanque/linea y es accesible/autorizado."],
        ["Etiqueta", "Codigo unico, producto, origen, fecha/hora, punto, responsable y sello."],
        ["Seguridad", "Permiso, EPP, control de atmosfera y reglas del sitio cuando apliquen."],
        ["Documentos", "Plan, instruccion, especificacion y cadena de custodia disponibles."],
    ], [1.35, 5.37])
    doc.add_heading("7. Observacion del muestreo", level=1)
    for item in [
        "Registrar quién ejecuta, quién observa y quién recibe la muestra.",
        "Documentar fecha, hora, zona horaria, punto, secuencia y condiciones visibles.",
        "Verificar homogeneizacion, purga o acondicionamiento solo conforme al metodo aplicable.",
        "Evitar contacto, contaminacion cruzada, exposicion o trasvasije no autorizado.",
        "Registrar cualquier desviacion sin intentar corregir una tecnica para la cual ACI no este autorizado.",
    ]:
        bullet(doc, item)
    doc.add_heading("8. Etiquetado, sellado y custodia", level=1)
    table(doc, ["Elemento", "Registro minimo"], [
        ["Codigo", "Identificador unico e irrepetible."],
        ["Sello", "Numero, condicion, aplicacion, apertura y responsable."],
        ["Transferencia", "Fecha/hora, entrega, recepcion, firmas o evidencia equivalente."],
        ["Transporte", "Medio, condiciones, embalaje, temperatura si aplica y eventos."],
        ["Recepcion", "Laboratorio/persona, fecha/hora, estado y discrepancias."],
        ["Retencion", "Ubicacion, plazo, acceso y disposicion final autorizada."],
    ], [1.35, 5.37])
    doc.add_heading("9. Revision de resultados", level=1)
    numbered(doc, "Identificar", "Relacionar certificado, muestra, lote, metodo, fecha y laboratorio.")
    numbered(doc, "Comprobar", "Verificar unidades, limites, resultados, observaciones y aprobaciones del certificado.")
    numbered(doc, "Comparar", "Usar exclusivamente la especificacion o criterio proporcionado y vigente.")
    numbered(doc, "Informar", "Distinguir resultado reportado por laboratorio de cualquier observacion de ACI.")
    numbered(doc, "Escalar", "Notificar resultado fuera de especificacion, muestra comprometida o certificado inconsistente.")
    doc.add_heading("10. Muestras no conformes", level=1)
    for item in ["Sello roto o no coincidente.", "Etiqueta incompleta o inconsistente.", "Recipiente dañado, sucio, incompatible o con fuga.", "Cadena de custodia incompleta.", "Volumen insuficiente o preservacion no demostrada.", "Resultado sin trazabilidad a la muestra o metodo."]:
        bullet(doc, item)
    callout(doc, "Accion", "No desechar ni reemplazar evidencia por iniciativa propia. Segregar cuando sea posible, fotografiar, registrar y solicitar instruccion escrita.", CAUTION)
    common_end(doc, ["ACI-F-003 Bitacora.", "ACI-F-004 Registro Fotografico.", "Cadena de custodia del cliente/laboratorio.", "Certificados de calidad.", "ACI-F-005 y ACI-F-006."], "Ejecutar una simulacion con tres muestras, una etiqueta incompleta y un sello discrepante. Verificar que el equipo preserve evidencia, escale y documente la limitacion.")
    return code, "ACI-PT-002-Muestreo-Calidad-Cadena-Custodia.docx", doc


def expediting_doc():
    code = "ACI-PT-003"
    doc = base_doc(code, "Expediting y Tiempos")
    cover(doc, code, "Procedimiento Tecnico", "Expediting y Control de Tiempos")
    common_start(doc, code, "Definir el metodo para controlar hitos, tiempos, restricciones, comunicaciones y causas reportadas durante operaciones petroleras, permitiendo alertas tempranas y una cronologia defendible.", "Aplica desde la programacion y arribo hasta la finalizacion, documentacion y salida, según el alcance establecido en la orden de servicio.")
    doc.add_heading("4. Principios", level=1)
    for item in ["Utilizar una zona horaria declarada y mantenerla en todos los registros.", "Diferenciar hora observada, hora reportada por tercero y hora calculada.", "Registrar hechos y causa reportada sin atribuir responsabilidad contractual o legal.", "Comunicar impacto potencial antes de que la demora se consolide.", "No modificar retrospectivamente la bitacora; agregar correccion trazable."]:
        bullet(doc, item)
    doc.add_heading("5. Plan de hitos", level=1)
    table(doc, ["Fase", "Hitos posibles"], [
        ["Programacion", "Nominacion, ETA, ventana, cambios, instrucciones, disponibilidad."],
        ["Arribo", "Llegada, fondeo, atraque, NOR u otro hito contractual reportado."],
        ["Preparacion", "Acceso, induccion, reunion, documentos, inspecciones, conexiones, readiness."],
        ["Transferencia", "Inicio, caudal, cambios de tanque, detenciones, reinicios, termino."],
        ["Medicion", "Lecturas iniciales/finales, muestreo, calculos, documentos."],
        ["Cierre", "Desconexion, documentos finales, salida, reporte preliminar."],
    ], [1.35, 5.37])
    doc.add_heading("6. Registro de eventos", level=1)
    for item in ["Fecha y hora exacta.", "Evento observado o comunicado.", "Fuente y persona que informa.", "Estado anterior y posterior.", "Causa reportada, sin convertirla en conclusion de ACI.", "Accion adoptada, destinatarios y respuesta.", "Evidencia relacionada."]:
        bullet(doc, item)
    doc.add_heading("7. Detenciones y demoras", level=1)
    numbered(doc, "Abrir", "Registrar inicio y condicion que impide o reduce la operacion.")
    numbered(doc, "Confirmar", "Solicitar causa y parte que la reporta, evitando suposiciones.")
    numbered(doc, "Monitorear", "Actualizar estado, impacto estimado y acciones de recuperacion.")
    numbered(doc, "Cerrar", "Registrar hora de reinicio, duracion y condicion al reanudar.")
    numbered(doc, "Reconciliar", "Comparar la bitacora con SOF, logs y comunicaciones disponibles.")
    doc.add_heading("8. Alertas", level=1)
    table(doc, ["Condicion", "Alerta recomendada"], [
        ["Hito vencido", "Informar atraso, causa reportada, impacto y siguiente actualizacion."],
        ["Detencion no prevista", "Comunicar inmediatamente cuando pueda afectar ventana, cantidad o costo."],
        ["Informacion contradictoria", "Presentar versiones y solicitar confirmacion escrita."],
        ["Cambio de secuencia", "Registrar autorizacion, motivo e impacto operacional."],
        ["Documento pendiente", "Identificar responsable y plazo antes del cierre."],
    ], [2.0, 4.72])
    doc.add_heading("9. Informe de tiempos", level=1)
    paragraph(doc, "El informe debe incluir cronologia, periodos de actividad e inactividad, causa reportada, fuente, comunicaciones, pendientes y limitaciones. Cualquier calculo de laytime, demurrage u obligación contractual solo se emitira si está expresamente contratado y revisado por personal competente.")
    common_end(doc, ["ACI-F-001 Orden de Servicio.", "ACI-F-003 Bitacora.", "Comunicaciones y documentos de terceros.", "ACI-F-005 Reporte Preliminar.", "ACI-F-006 Informe Final."], "Simular una operacion con cambio de ETA, detencion, causa contradictoria y reinicio. Verificar zona horaria, fuente, alertas y reconciliacion de cronologia.")
    return code, "ACI-PT-003-Expediting-Control-Tiempos.docx", doc


def evidence_doc():
    code = "ACI-PT-004"
    doc = base_doc(code, "Evidencia y Reportabilidad")
    cover(doc, code, "Procedimiento Tecnico", "Evidencia, Fotografias y Reportabilidad")
    common_start(doc, code, "Establecer cómo ACI crea, identifica, preserva, revisa y entrega evidencia operacional e informes trazables.", "Aplica a bitacoras, fotografias, videos autorizados, mediciones, certificados, comunicaciones, calculos, reportes y expedientes digitales del servicio.")
    doc.add_heading("4. Clasificacion de evidencia", level=1)
    table(doc, ["Tipo", "Ejemplos", "Control"], [
        ["Primaria observada", "Medicion presenciada, fotografia propia, evento registrado.", "Autor, fecha/hora, lugar y contexto."],
        ["Recibida", "Certificado, log, email, documento de nave/terminal.", "Fuente, fecha/hora de recepcion y version."],
        ["Calculada", "Reconciliacion, duracion, diferencia.", "Datos fuente, formula, unidad, autor y revisor."],
        ["Declarativa", "Explicacion verbal o escrita de tercero.", "Nombre, rol, momento y texto objetivo."],
    ], [1.45, 2.75, 2.52])
    doc.add_heading("5. Expediente digital", level=1)
    paragraph(doc, "Cada orden debe tener una carpeta unica. Estructura minima recomendada:")
    for item in ["01-Instrucciones", "02-Operacional", "03-Mediciones", "04-Calidad", "05-Fotografias", "06-Comunicaciones", "07-Reportes", "08-Cierre"]:
        bullet(doc, item)
    doc.add_heading("6. Nomenclatura", level=1)
    callout(doc, "Formato", "OS-AAAA-NNN_Tipo_FechaHora_Descripcion_Version.ext. No usar nombres ambiguos como final2, nuevo o foto1 sin relacion con el registro.")
    doc.add_heading("7. Registro fotografico", level=1)
    for item in [
        "Fotografiar solo cuando sea seguro, permitido y relevante para el servicio.",
        "Mantener archivo original y registrar numero, autor, fecha/hora, ubicación, objeto y descripcion.",
        "Evitar personas, pantallas, credenciales o información innecesaria.",
        "No editar el original. Cualquier copia marcada o recortada debe identificarse como derivada.",
        "No afirmar escala, medida o condicion no demostrable por la imagen.",
    ]:
        bullet(doc, item)
    doc.add_heading("8. Integridad y respaldo", level=1)
    table(doc, ["Control", "Aplicacion"], [
        ["Original", "Conservar el archivo recibido o creado sin sobreescribirlo."],
        ["Version", "Emitir nueva version y registrar cambio, autor, fecha y motivo."],
        ["Respaldo", "Copiar al repositorio autorizado tan pronto como sea posible."],
        ["Acceso", "Restringir por orden, cliente y función."],
        ["Transmision", "Usar canales autorizados y destinatarios confirmados."],
        ["Retencion", "Aplicar contrato, ley y matriz aprobada; no eliminar unilateralmente."],
    ], [1.35, 5.37])
    doc.add_heading("9. Reporte preliminar", level=1)
    for item in ["Estado y alcance cubierto.", "Hechos relevantes y cronologia.", "Cantidades o resultados claramente marcados como preliminares.", "Desviaciones y alertas emitidas.", "Documentos o datos pendientes.", "Proximos pasos y plazo estimado de informe final."]:
        bullet(doc, item)
    doc.add_heading("10. Informe final", level=1)
    table(doc, ["Seccion", "Contenido"], [
        ["Resumen", "Resultado ejecutivo, principales hechos y limitaciones."],
        ["Alcance", "Instruccion, actividades realizadas, exclusiones y restricciones."],
        ["Metodologia", "Fuentes, metodos, unidades y criterios utilizados."],
        ["Cronologia", "Eventos relevantes y comunicaciones."],
        ["Resultados", "Cantidad, calidad, tiempos y documentos según alcance."],
        ["Hallazgos", "Hecho, evidencia, impacto potencial y acción."],
        ["Conclusiones", "Conclusiones técnicas prudentes y sustentadas."],
        ["Anexos", "Documentos, fotos, calculos y referencias."],
    ], [1.35, 5.37])
    doc.add_heading("11. Revision y emision", level=1)
    for item in ["Correspondencia con orden de servicio.", "Coherencia de cronologia y cifras.", "Evidencia identificable para hallazgos.", "Diferenciacion de datos propios y de terceros.", "Lenguaje técnico, objetivo y no acusatorio.", "Control de versión, aprobacion y destinatarios.", "Eliminacion de comentarios internos y metadatos innecesarios."]:
        bullet(doc, item)
    common_end(doc, ["ACI-F-003 Bitacora.", "ACI-F-004 Registro Fotografico.", "ACI-F-005 Reporte Preliminar.", "ACI-F-006 Informe Final.", "ACI-F-101 Listado Maestro.", "ACI-F-107 Matriz de Retencion."], "Crear un expediente simulado, introducir un archivo duplicado y una fotografia sin contexto, y verificar nomenclatura, correccion, revision, respaldo y entrega.")
    return code, "ACI-PT-004-Evidencia-Reportabilidad.docx", doc


def claims_doc():
    code = "ACI-PT-005"
    doc = base_doc(code, "Diferencias y Reclamos")
    cover(doc, code, "Procedimiento Tecnico", "Investigacion de Diferencias y Soporte a Reclamos")
    common_start(doc, code, "Definir el proceso para preservar evidencia, analizar diferencias, construir una cronologia y emitir soporte tecnico para decisiones o reclamos, sin asumir funciones legales ni atribuir responsabilidad fuera de la evidencia.", "Aplica a diferencias de cantidad, calidad, tiempo, documentos, sellos, muestras y eventos operacionales relacionados con servicios ACI.")
    doc.add_heading("4. Limites", level=1)
    callout(doc, "Limite", "ACI entrega analisis tecnico. No emite opinion legal, no determina responsabilidad contractual y no representa al cliente ante terceros salvo autorizacion expresa.", CAUTION)
    doc.add_heading("5. Activacion", level=1)
    for item in ["Diferencia superior al criterio contractual o interno aprobado.", "Resultado de calidad fuera de especificacion o muestra comprometida.", "Demora relevante o cronologia controvertida.", "Documento, sello o medicion inconsistente.", "Solicitud formal del cliente para preservar o analizar evidencia."]:
        bullet(doc, item)
    doc.add_heading("6. Preservacion", level=1)
    numbered(doc, "Congelar", "Conservar originales, versiones, correos, fotografias, calculos y registros sin sobreescritura.")
    numbered(doc, "Identificar", "Registrar fuente, fecha/hora, custodio y relación con la operación.")
    numbered(doc, "Respaldar", "Copiar al repositorio restringido y controlar acceso.")
    numbered(doc, "Solicitar", "Identificar documentos faltantes y pedirlos por canal trazable.")
    numbered(doc, "Registrar", "Documentar cualquier perdida, alteracion, demora o limitacion de evidencia.")
    doc.add_heading("7. Plan de investigacion", level=1)
    table(doc, ["Pregunta", "Control"], [
        ["Que ocurrió", "Definir la diferencia o controversia con unidad, periodo y alcance."],
        ["Cuando", "Construir cronologia con zona horaria y fuentes."],
        ["Donde", "Aislar tanque, compartimiento, linea, punto o etapa."],
        ["Con qué datos", "Separar observados, calculados, recibidos y declarativos."],
        ["Que falta", "Crear lista de evidencia pendiente y responsable."],
        ["Que hipótesis", "Registrar explicaciones posibles sin convertirlas en hechos."],
    ], [1.55, 5.17])
    doc.add_heading("8. Analisis", level=1)
    for item in [
        "Recalcular con formulas visibles y segundo revisor.",
        "Comparar documentos independientes y buscar contradicciones.",
        "Aislar el momento en que aparece la diferencia.",
        "Evaluar causas alternativas y evidencia que las contradice.",
        "Distinguir correlacion temporal de causa demostrada.",
        "Cuantificar incertidumbre y limitaciones cuando sea posible.",
    ]:
        bullet(doc, item)
    doc.add_heading("9. Comunicaciones y protestas", level=1)
    for item in ["No firmar admision de responsabilidad.", "No alterar documentos de terceros.", "Registrar documentos firmados con reserva, rechazados o recibidos.", "Utilizar lenguaje factual y solicitar instrucciones para comunicaciones externas.", "Escalar inmediatamente amenazas legales, regulatorias, ambientales o reputacionales."]:
        bullet(doc, item)
    doc.add_heading("10. Informe tecnico", level=1)
    table(doc, ["Seccion", "Contenido"], [
        ["Mandato", "Solicitud, alcance, preguntas y limitaciones."],
        ["Evidencia", "Inventario, fuentes, integridad y datos faltantes."],
        ["Cronologia", "Eventos verificados y fuente de cada uno."],
        ["Calculos", "Datos, formulas, resultados y revisión."],
        ["Hallazgos", "Hechos sustentados y observaciones."],
        ["Hipotesis", "Explicaciones posibles y nivel de soporte."],
        ["Conclusiones", "Respuesta técnica prudente a las preguntas."],
        ["Anexos", "Documentos, tablas, fotos y referencias."],
    ], [1.35, 5.37])
    common_end(doc, ["ACI-F-003 Bitacora.", "ACI-F-004 Registro Fotografico.", "ACI-F-006 Informe Final.", "Inventario de evidencia.", "Calculos revisados.", "Registro de comunicaciones."], "Simular una diferencia con dos fuentes contradictorias, un documento faltante y una hipótesis plausible no demostrada. Verificar que el informe mantenga la distinción entre hecho, cálculo e hipótesis.")
    return code, "ACI-PT-005-Diferencias-Soporte-Reclamos.docx", doc


def safety_doc():
    code = "ACI-HSE-001"
    doc = base_doc(code, "Seguridad Operacional")
    cover(doc, code, "Manual HSE", "Seguridad Operacional y Respuesta a Emergencias")
    common_start(doc, code, "Establecer requisitos mínimos de seguridad, salud, ambiente y respuesta a emergencias para personal ACI y proveedores durante servicios en terreno.", "Aplica a terminales, buques, plantas, patios, rutas, oficinas temporales y desplazamientos relacionados con servicios ACI. Siempre prevalecen la ley y las reglas del sitio cuando sean más exigentes.")
    doc.add_heading("4. Responsabilidades y autoridad", level=1)
    table(doc, ["Rol", "Responsabilidad"], [
        ["Gerencia", "Recursos, política, seguros, investigación y decisiones de alto riesgo."],
        ["Coordinador", "Requisitos del sitio, contactos, aptitud, viajes y comunicaciones."],
        ["Inspector", "Cumplir reglas, evaluar condiciones, usar EPP y detener su actividad insegura."],
        ["Proveedor", "Demostrar competencia, aptitud, equipos y cumplimiento aplicable."],
    ], [1.45, 5.27])
    callout(doc, "Autoridad para detener", "Toda persona ACI puede suspender su propia actividad ante peligro no controlado, falta de autorización, EPP inadecuado, fatiga incapacitante o instruccion insegura. Debe retirarse a zona segura y comunicar.", CAUTION)
    doc.add_heading("5. Preparacion del servicio", level=1)
    for item in ["Requisitos de acceso, induccion, permisos y documentación.", "Peligros del producto y hojas de seguridad disponibles.", "EPP mínimo y específico del sitio.", "Aptitud médica, fatiga, jornada, descanso y trabajo en solitario.", "Transporte, alojamiento, clima y seguridad personal.", "Contactos de emergencia, punto de reunión y medios de comunicación.", "Restricciones de fotografias, teléfonos y equipos electrónicos."]:
        bullet(doc, item)
    doc.add_heading("6. Evaluacion de riesgo en terreno", level=1)
    numbered(doc, "Detener", "Antes de iniciar, observar el área y confirmar autorización.")
    numbered(doc, "Identificar", "Energía, tráfico, caídas, atmósfera, químicos, presión, superficies, clima y simultaneidad.")
    numbered(doc, "Controlar", "Aplicar jerarquía de controles y reglas del sitio.")
    numbered(doc, "Confirmar", "Verificar que controles, EPP, comunicación y salida estén disponibles.")
    numbered(doc, "Reevaluar", "Repetir ante cambio de tarea, ubicación, clima, producto o condición.")
    doc.add_heading("7. Peligros mínimos a considerar", level=1)
    table(doc, ["Peligro", "Controles mínimos"], [
        ["Atmósferas inflamables/tóxicas", "Zonas autorizadas, detección y permisos del sitio, equipos aprobados, no ingresar a espacios confinados."],
        ["Presión y transferencia", "Mantener distancia, no operar válvulas, reconocer mangueras/lineas y puntos de liberación."],
        ["Caídas y acceso", "Tres puntos de contacto, pasarelas autorizadas, iluminación y protección requerida."],
        ["Vehículos/equipos móviles", "Rutas peatonales, alta visibilidad, contacto visual y zonas segregadas."],
        ["Producto químico", "SDS, guantes/antiparras adecuados, higiene y respuesta a exposición."],
        ["Marítimo", "Reglas de embarque, chaleco cuando corresponda, clima, mareas y transferencia segura."],
        ["Fatiga", "Límites de jornada, pausas, relevo, transporte seguro y declaración temprana."],
        ["Clima", "Calor, frío, tormenta, viento, lluvia y suspensión según reglas del sitio."],
    ], [1.65, 5.07])
    doc.add_heading("8. EPP", level=1)
    paragraph(doc, "El EPP se define por evaluación de riesgo y sitio. Como base orientativa, puede incluir casco, lentes, ropa de manga larga/retardante cuando aplique, calzado de seguridad, guantes adecuados y alta visibilidad. Protección respiratoria, auditiva, contra caídas o flotación requiere selección, entrenamiento y autorización específica.")
    doc.add_heading("9. Prohibiciones", level=1)
    for item in ["Ingresar a espacios confinados o zonas no autorizadas.", "Operar válvulas, bombas, sellos o equipos del cliente sin autorización y competencia.", "Anular controles o permisos.", "Trabajar bajo efectos de alcohol, drogas o fatiga incapacitante.", "Usar dispositivos no autorizados en áreas clasificadas.", "Tomar muestras o exponerse al producto sin método, EPP y autorización."]:
        bullet(doc, item)
    doc.add_heading("10. Emergencias", level=1)
    numbered(doc, "Proteger", "No convertirse en víctima; alejarse a zona segura.")
    numbered(doc, "Alertar", "Activar alarma o informar por el canal del sitio.")
    numbered(doc, "Evacuar", "Seguir rutas, punto de reunión y conteo.")
    numbered(doc, "Informar", "Comunicar a Coordinador ACI tan pronto sea seguro.")
    numbered(doc, "Preservar", "Registrar hechos y evidencia solo cuando esté autorizado y sea seguro.")
    doc.add_heading("11. Incidentes y condiciones inseguras", level=1)
    for item in ["Reportar inmediatamente lesión, exposición, derrame, incendio, casi incidente o daño.", "Obtener atención médica y cumplir reporte del sitio.", "No emitir declaraciones externas ni admitir responsabilidad.", "Preservar registros y colaborar con investigación autorizada.", "No reiniciar actividad hasta confirmar controles y autorización."]:
        bullet(doc, item)
    doc.add_heading("12. Viajes y trabajo aislado", level=1)
    for item in ["Itinerario y contacto de seguimiento.", "Medio de transporte y conductor adecuados.", "Registro de llegada/salida en ubicaciones remotas.", "Evitar conducir fatigado o después de jornadas extendidas.", "Medio de comunicación alternativo y criterio de escalamiento.", "Evaluación de seguridad personal y condiciones locales."]:
        bullet(doc, item)
    common_end(doc, ["Registro de induccion.", "Evaluacion de riesgo/LMRA del sitio cuando aplique.", "Registro de incidente o condición insegura.", "ACI-F-003 Bitacora.", "Registro de capacitación HSE."], "Realizar un ejercicio de escritorio con alarma, evacuación, persona no localizada y comunicación al cliente. Verificar roles, contactos, preservación de evidencia y retorno seguro.")
    return code, "ACI-HSE-001-Seguridad-Emergencias.docx", doc


def training_doc():
    code = "ACI-PC-001"
    doc = base_doc(code, "Capacitacion y Habilitacion")
    cover(doc, code, "Programa Corporativo", "Capacitacion y Habilitacion Tecnica")
    common_start(doc, code, "Definir el programa mediante el cual ACI forma, evalúa, supervisa, habilita y reevalúa al personal que ejecuta o revisa servicios de loss control.", "Aplica a inspectores, loss controllers, coordinadores, revisores técnicos y proveedores críticos que actúan en nombre de ACI.")
    doc.add_heading("4. Niveles de habilitacion", level=1)
    table(doc, ["Nivel", "Criterio", "Autorizacion"], [
        ["En induccion", "Documentos corporativos y HSE en curso.", "Sin trabajo técnico autónomo."],
        ["En formación", "Teoría aprobada; práctica supervisada pendiente o parcial.", "Participa acompañado; no firma informes."],
        ["Habilitado", "Competencias mínimas, práctica y evaluación satisfactorias.", "Ejecuta alcances definidos."],
        ["Senior", "Experiencia y criterio demostrados en operaciones complejas.", "Lidera y supervisa."],
        ["Revisor técnico", "Dominio metodológico, independencia y evaluación específica.", "Aprueba técnicamente reportes."],
    ], [1.15, 3.45, 2.12])
    doc.add_heading("5. Malla formativa", level=1)
    table(doc, ["Modulo", "Contenido mínimo", "Evaluacion"], [
        ["M1 ACI e independencia", "Alcance, ética, confidencialidad, conflicto y escalamiento.", "Prueba y declaración."],
        ["M2 HSE", "Peligros, EPP, autoridad para detener, emergencia, fatiga y viajes.", "Prueba y ejercicio."],
        ["M3 Operación petrolera", "Carga/descarga, terminal, nave, tanques, documentos y secuencia.", "Caso técnico."],
        ["M4 Cantidad", "Medición, unidades, factores, reconciliación y revisión.", "Cálculo controlado."],
        ["M5 Calidad", "Muestreo, etiquetas, sellos, custodia y certificados.", "Simulación."],
        ["M6 Expediting", "Hitos, zona horaria, demoras, fuentes y alertas.", "Cronología simulada."],
        ["M7 Evidencia", "Bitácora, fotos, nomenclatura, respaldo e integridad.", "Expediente simulado."],
        ["M8 Reportes", "Preliminar, informe final, lenguaje y revisión.", "Informe evaluado."],
        ["M9 Reclamos", "Preservación, hipótesis, cálculos y límites legales.", "Caso de diferencia."],
    ], [1.45, 3.85, 1.42])
    doc.add_heading("6. Evaluacion", level=1)
    paragraph(doc, "La habilitación exige evidencia combinada. Una prueba escrita por sí sola no demuestra competencia operacional.")
    table(doc, ["Componente", "Criterio inicial propuesto"], [
        ["Conocimiento", "Resultado >= 80% en módulos aplicables."],
        ["Caso práctico", "Sin errores críticos y >= 80% de criterios."],
        ["HSE", "100% de puntos críticos y conducta segura."],
        ["Operación supervisada", "Evaluación satisfactoria del supervisor."],
        ["Reporte", "Aprobado por revisor técnico sin reproceso mayor."],
        ["Conducta", "Independencia, comunicación y trazabilidad satisfactorias."],
    ], [2.0, 4.72])
    paragraph(doc, "Los porcentajes son criterios iniciales y deben ser aprobados por Gerencia y Calidad antes de su vigencia.", italic=True, color=STEEL)
    doc.add_heading("7. Practica supervisada", level=1)
    for item in ["Revisión previa del alcance y riesgos.", "Participación en operación real o simulación de fidelidad suficiente.", "Observación directa de preparación, registro, comunicación y cierre.", "Retroalimentación documentada.", "Repetición de actividades insuficientes.", "Autorización formal por alcance, no genérica."]:
        bullet(doc, item)
    doc.add_heading("8. Habilitacion por competencia", level=1)
    table(doc, ["Competencia", "Evidencia"], [
        ["Operaciones", "Caso y servicio supervisado."],
        ["Cantidad", "Ejercicio de medición/reconciliación revisado."],
        ["Calidad", "Simulación de muestreo/custodia."],
        ["Expediting", "Cronología y alertas correctas."],
        ["Evidencia", "Expediente completo y trazable."],
        ["Reportes", "Preliminar y final aprobados."],
        ["Seguridad", "Inducción, evaluación y desempeño seguro."],
    ], [1.55, 5.17])
    doc.add_heading("9. Reevaluacion", level=1)
    for item in ["Anual o según riesgo y frecuencia de actividad.", "Después de ausencia prolongada en una competencia.", "Tras error crítico, incidente, reclamo o no conformidad.", "Cuando cambie método, producto, equipo, requisito o jurisdicción.", "Antes de asumir revisión técnica o liderazgo de operación compleja."]:
        bullet(doc, item)
    doc.add_heading("10. Suspension o restriccion", level=1)
    paragraph(doc, "Operaciones y Calidad pueden suspender o limitar una habilitación por vencimiento, desempeño insuficiente, incumplimiento HSE, alteración de registros, conflicto de interés no declarado o falta de actividad reciente. La restitución requiere plan y evidencia de cierre.")
    doc.add_heading("11. Revisor tecnico", level=1)
    for item in ["Competencia técnica superior o equivalente al alcance revisado.", "Independencia respecto del trabajo cuando el riesgo lo requiera.", "Capacidad de reconstruir cálculos y desafiar conclusiones.", "Conocimiento de control documental y confidencialidad.", "Registro de revisiones y desempeño satisfactorio."]:
        bullet(doc, item)
    common_end(doc, ["ACI-F-201 Perfil de Cargo.", "ACI-F-202 Matriz de Competencia.", "ACI-F-203 Registro de Capacitación.", "ACI-F-204 Evaluación de Desempeño.", "Evaluaciones y casos prácticos.", "Registro de habilitación/suspensión."], "Aplicar el programa a una persona ficticia, incluyendo brecha, formación, caso práctico, operación supervisada, evaluación y decisión de habilitación por competencia.")
    return code, "ACI-PC-001-Capacitacion-Habilitacion.docx", doc


def index_doc(items):
    code = "ACI-OP-000"
    doc = base_doc(code, "Indice Operacional")
    cover(doc, code, "Indice Maestro", "Documentacion Operacional y Tecnica")
    doc.add_heading("1. Objeto", level=1)
    paragraph(doc, "Identificar la documentación operacional desarrollada para ACI LATAM y su estado de implementación.")
    rows = [["ACI-MO-001", "Manual Operacional de Loss Control y Expediting", "Para aprobación", "Marco general"]]
    for item_code, filename, _ in items:
        title = filename.replace(".docx", "").replace(item_code + "-", "").replace("-", " ")
        rows.append([item_code, title, "Para aprobación", "Procedimiento / programa"])
    rows.extend([
        ["ACI-PQ-001", "Paquete de Formatos Operacionales", "Para aprobación", "Excel ACI-F-001 a ACI-F-006"],
        ["ACI-PQ-002", "Control Documental y Competencia", "Para aprobación", "Excel ACI-F-101 y ACI-F-201 a 204"],
    ])
    table(doc, ["Codigo", "Documento", "Estado", "Funcion"], rows, [1.1, 3.35, 1.25, 1.02])
    doc.add_heading("2. Condicion de vigencia", level=1)
    paragraph(doc, "Ningún documento se considera vigente hasta contar con responsable nominal, aprobación, capacitación, prueba piloto y registro en el listado maestro de documentos.")
    doc.add_heading("3. Orden de implementacion", level=1)
    for text in ["Aprobar ACI-MO-001 y ACI-HSE-001.", "Designar responsables y completar matrices de competencia.", "Capacitar con ACI-PC-001.", "Simular los procedimientos ACI-PT-001 a ACI-PT-005.", "Ajustar formatos y emitir versiones vigentes.", "Ejecutar operaciones piloto supervisadas."]:
        bullet(doc, text)
    return code, "ACI-OP-000-Indice-Operacional.docx", doc


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    generated = []
    for builder in [quantity_doc, quality_doc, expediting_doc, evidence_doc, claims_doc, safety_doc, training_doc]:
        code, filename, doc = builder()
        path = OUT / filename
        doc.save(path)
        generated.append((code, filename, path))
    _, filename, index = index_doc(generated)
    index_path = OUT / filename
    index.save(index_path)
    print(index_path)
    for _, _, path in generated:
        print(path)


if __name__ == "__main__":
    build()
