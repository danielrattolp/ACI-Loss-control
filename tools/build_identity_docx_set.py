from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
IDENTITY = ROOT / "docs" / "identidad"
LOGO = ROOT / "assets" / "aci-logo-black.jpeg"

INK = "0B0F12"
WHITE = "FFFFFF"
AMBER = "C88C3A"
SEA = "1F5867"
STEEL = "5D788A"
PAPER = "F6F7F4"
LINE = "D7DDE1"

DOCS = [
    (
        "brand-book.md",
        "ACI-Brand-Book.docx",
        "Brand Book",
        "Identidad estrategica, verbal y visual de ACI Loss Control Experts.",
    ),
    (
        "arquitectura-de-marca.md",
        "ACI-Arquitectura-de-Marca.docx",
        "Arquitectura de Marca",
        "Estructura de marca, lineas de servicio y reglas de crecimiento.",
    ),
    (
        "lineamientos-visuales.md",
        "ACI-Lineamientos-Visuales.docx",
        "Lineamientos Visuales",
        "Uso de logo, paleta, fotografia, iconografia y composicion.",
    ),
    (
        "mensajes-comerciales.md",
        "ACI-Mensajes-Comerciales.docx",
        "Mensajes Comerciales",
        "Claim, pitch, descripciones, beneficios y argumentos comerciales.",
    ),
]


def set_font(run, size=None, color=None, bold=None, italic=None, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("ACI Loss Control Experts")
    set_font(r, size=8, color="707A82")
    return doc


def add_cover(doc, title, subtitle):
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover.autofit = False
    cover.columns[0].width = Inches(6.5)
    set_table_borders(cover, color=INK, size="0")
    cell = cover.cell(0, 0)
    set_cell_shading(cell, INK)
    set_cell_margins(cell, top=260, start=260, bottom=260, end=260)

    p_logo = cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p_logo.add_run().add_picture(str(LOGO), width=Inches(3.35))

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_font(r, size=24, color=WHITE, bold=True)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(subtitle)
    set_font(r2, size=11.5, color="DCE3E7")

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Chile | Ecuador | Colombia")
    set_font(r3, size=10, color=AMBER, bold=True)

    doc.add_page_break()


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        size, color = 16, SEA
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        size, color = 13, SEA
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        size, color = 11.5, STEEL
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=True)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, size=10.5, color="25313A")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, size=10.5, color="25313A")


def parse_table(lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        parts = [part.strip() for part in stripped.strip("|").split("|")]
        if all(set(part.replace(" ", "")) <= {"-", ":"} for part in parts):
            continue
        rows.append(parts)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    width = 6.5 / cols

    for idx, value in enumerate(rows[0]):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(width)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, INK)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=9, color=WHITE, bold=True)

    for row in rows[1:]:
        cells = table.add_row().cells
        for idx in range(cols):
            cell = cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if idx == 0:
                set_cell_shading(cell, PAPER)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(row[idx] if idx < len(row) else "")
            set_font(r, size=9, color="303A40", bold=(idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_callout_for_claim(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    set_table_borders(table, color="D9C29A", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7E8")
    set_cell_margins(cell, top=150, start=190, bottom=150, end=190)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=14, color=INK, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_from_markdown(md_path, out_path, title, subtitle):
    doc = setup_document()
    add_cover(doc, title, subtitle)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    table_buffer = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, parse_table(table_buffer))
            table_buffer = []

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_table()
            continue
        if stripped.startswith("|"):
            table_buffer.append(stripped)
            continue
        flush_table()
        if stripped.startswith("# "):
            text = stripped[2:].strip()
            if text != title:
                add_heading(doc, text, 1)
        elif stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 1)
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 2)
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:].strip())
        elif stripped and all(ch == "-" for ch in stripped):
            continue
        else:
            if title == "Mensajes Comerciales" and "Control tecnico donde cada barril importa" in stripped:
                add_callout_for_claim(doc, stripped)
            else:
                add_paragraph(doc, stripped)
    flush_table()
    doc.save(out_path)
    return out_path


def main():
    outputs = []
    for source, target, title, subtitle in DOCS:
        outputs.append(build_from_markdown(IDENTITY / source, IDENTITY / target, title, subtitle))
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
