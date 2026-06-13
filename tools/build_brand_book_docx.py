from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "identidad" / "ACI-Manual-Identidad-Corporativa.docx"
LOGO = ROOT / "assets" / "aci-logo-black.jpeg"
HERO = ROOT / "assets" / "oil-tanker-hero.png"

INK = "0B0F12"
WHITE = "FFFFFF"
AMBER = "C88C3A"
STEEL = "5D788A"
SEA = "1F5867"
PAPER = "F6F7F4"
LINE = "D7DDE1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_font(run, size=None, color=None, bold=None, italic=None, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", size=10.5, color="1F2933", bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        size, color = 16, SEA
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        size, color = 13, SEA
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        size, color = 11.5, STEEL
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if p.runs:
        p.runs[0].text = text
        set_font(p.runs[0], size=10.5, color="1F2933")
    else:
        r = p.add_run(text)
        set_font(r, size=10.5, color="1F2933")
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, size=10.5, color="1F2933")
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    set_table_borders(table, color="D9C29A", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7E8")
    set_cell_margins(cell, top=170, start=190, bottom=160, end=190)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, size=10.5, color=INK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, size=10, color="303A40")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_kv_table(doc, rows, widths=(1.85, 4.65)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=LINE, size="6")
    for row in table.rows:
        row.cells[0].width = Inches(widths[0])
        row.cells[1].width = Inches(widths[1])
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        for c in (c0, c1):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(c, top=95, start=140, bottom=95, end=140)
        set_cell_shading(c0, PAPER)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_font(r0, size=9.5, color=INK, bold=True)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_font(r1, size=9.5, color="303A40")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=LINE, size="6")
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_shading(cell, INK)
        set_cell_margins(cell, top=110, start=120, bottom=110, end=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=9, color=WHITE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_margins(cells[i], top=100, start=120, bottom=100, end=120)
            if i == 0:
                set_cell_shading(cells[i], PAPER)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=9, color="303A40", bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_section_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("ACI Loss Control Experts | Manual de Identidad Corporativa")
    set_font(r, size=8, color="707A82")


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)
        s.header_distance = Inches(0.35)
        s.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(10.5)
    add_section_footer(section)
    return doc


def cover(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    set_table_borders(table, color=INK, size="0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, INK)
    set_cell_margins(cell, top=260, start=260, bottom=260, end=260)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(3.8))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(22)
    p2.paragraph_format.space_after = Pt(4)
    r = p2.add_run("Manual de Identidad Corporativa")
    set_font(r, size=24, color=WHITE, bold=True)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Loss control petrolero, expediting y control documental")
    set_font(r3, size=12.5, color="DCE3E7")
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(14)
    r4 = p4.add_run("Chile | Ecuador | Colombia")
    set_font(r4, size=10.5, color=AMBER, bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    if HERO.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(str(HERO), width=Inches(6.5))
    doc.add_page_break()


def build():
    doc = setup_doc()
    cover(doc)

    add_heading(doc, "Indice", 1)
    for item in [
        "1. Esencia de marca",
        "2. Mision, vision y proposito",
        "3. Valores",
        "4. Personalidad y tono de voz",
        "5. Mensajes comerciales",
        "6. Oferta de servicios",
        "7. Audiencias",
        "8. Identidad visual",
        "9. Arquitectura de marca",
        "10. Aplicaciones iniciales",
    ]:
        add_para(doc, item, after=2)
    doc.add_page_break()

    add_heading(doc, "1. Esencia de marca", 1)
    add_kv_table(
        doc,
        [
            ("Nombre", "ACI Loss Control Experts."),
            ("Categoria", "Empresa especializada en loss control petrolero, expediting, reconciliacion de cantidades, control documental e inspeccion operacional de mercancias petroleras."),
            ("Idea central", "ACI protege los intereses del cliente en operaciones petroleras criticas, entregando presencia tecnica, evidencia verificable y reportabilidad clara desde la programacion hasta el informe final."),
            ("Posicionamiento", "Firma regional de loss control petrolero para operaciones de alto estandar en Chile, Ecuador y Colombia, con cobertura nacional, metodologia documentada y foco en trazabilidad."),
        ],
    )
    add_callout(
        doc,
        "Promesa de marca",
        "Donde existen diferencias de cantidad, calidad, tiempo o documentacion, ACI entrega control tecnico independiente para reducir perdidas, controversias y decisiones tomadas sin evidencia.",
    )

    add_heading(doc, "2. Mision, vision y proposito", 1)
    add_kv_table(
        doc,
        [
            ("Mision", "Proteger los intereses operacionales y economicos de nuestros clientes mediante servicios independientes de loss control, expediting, inspeccion y control documental de mercancias petroleras, asegurando trazabilidad, oportunidad y rigor tecnico."),
            ("Vision", "Ser la empresa latinoamericana de referencia en loss control petrolero para operaciones de alto requisito, reconocida por cobertura regional, confiabilidad tecnica y cultura de mejora continua."),
            ("Proposito", "Dar certeza a operaciones criticas, reduciendo perdidas, disputas y riesgos mediante presencia experta, evidencia verificable y decisiones informadas en tiempo real."),
        ],
        widths=(1.4, 5.1),
    )

    add_heading(doc, "3. Valores", 1)
    add_matrix(
        doc,
        ["Valor", "Definicion operacional"],
        [
            ("Independencia", "Actuamos con criterio tecnico, imparcialidad y transparencia. Nuestra credibilidad depende de informar lo observado, no lo conveniente."),
            ("Trazabilidad", "Todo hallazgo relevante debe tener respaldo: documento, medicion, fotografia, comunicacion, bitacora o calculo verificable."),
            ("Oportunidad", "En loss control, llegar tarde puede ser equivalente a no llegar. Priorizamos alertas tempranas y reportes preliminares."),
            ("Rigor tecnico", "Trabajamos con procedimientos, checklists, criterios de revision y control de evidencias."),
            ("Confidencialidad", "Protegemos informacion comercial, operacional y documental de cada cliente, operacion y contraparte involucrada."),
            ("Mejora continua", "Cada operacion debe dejar aprendizaje: desviaciones, oportunidades, registros y ajustes al sistema de trabajo."),
        ],
        widths=(1.55, 4.95),
    )

    add_heading(doc, "4. Personalidad y tono de voz", 1)
    add_para(doc, "ACI debe sentirse tecnica, seria, directa, regional, experta y operacional. Habla como un supervisor senior que conoce terreno, entiende el negocio y sabe documentar lo importante.")
    add_heading(doc, "Como hablamos", 2)
    for item in [
        "Con frases claras y concretas.",
        "Con vocabulario tecnico cuando aporta precision.",
        "Con foco en evidencia, tiempos, cantidades, calidad y documentos.",
        "Con seguridad, sin prometer resultados imposibles.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Como no hablamos", 2)
    for item in [
        "No usamos lenguaje grandilocuente sin respaldo.",
        "No declaramos certificaciones o acreditaciones no obtenidas.",
        "No prometemos eliminar todas las perdidas.",
        "No usamos claims genericos como lideres absolutos sin prueba.",
        "No nombramos clientes especificos sin autorizacion.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. Mensajes comerciales", 1)
    add_callout(doc, "Claim recomendado", "Control tecnico donde cada barril importa.")
    add_kv_table(
        doc,
        [
            ("Descripcion de una linea", "Loss control petrolero, expediting y control documental para operaciones criticas en Chile, Ecuador y Colombia."),
            ("Descripcion corta", "ACI protege los intereses del cliente durante operaciones petroleras criticas mediante presencia tecnica en terreno, reconciliacion de cantidades, control de calidad, control documental y reportabilidad trazable."),
            ("Elevator pitch", "Somos una empresa de loss control petrolero con cobertura en Chile, Ecuador y Colombia. Representamos tecnicamente los intereses del cliente durante operaciones de combustibles, controlando cantidad, calidad, tiempos, documentos y eventos relevantes."),
        ],
    )
    add_heading(doc, "Beneficios para el cliente", 2)
    for item in [
        "Menor exposicion a diferencias no explicadas.",
        "Mayor trazabilidad documental.",
        "Alertas tempranas durante la operacion.",
        "Evidencia tecnica para reclamos o cierres.",
        "Control externo sobre puntos criticos.",
        "Informacion clara para decidir.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Oferta de servicios", 1)
    add_matrix(
        doc,
        ["Servicio", "Alcance"],
        [
            ("Loss control y expediting", "Asistencia en carga, descarga y transferencia, control de tiempos, seguimiento de eventos y alertas tempranas."),
            ("Reconciliacion nave/terminal/planta", "Comparacion de cantidades, documentos, mediciones y diferencias operacionales entre las partes involucradas."),
            ("Control de cantidad y calidad", "Supervision de mediciones, muestreo, temperatura, aforo, certificados y documentacion de laboratorio."),
            ("Asistencia en carga y descarga", "Presencia tecnica durante la operacion para documentar eventos, desviaciones y condiciones relevantes."),
            ("Soporte a reclamos", "Consolidacion de evidencias, lineas de tiempo, documentos criticos, fotografias, calculos de diferencia y respaldo tecnico."),
            ("Informes y trazabilidad", "Reportes preliminares, bitacoras, resumen de operacion, dossier documental e informe final."),
        ],
        widths=(2.15, 4.35),
    )

    add_heading(doc, "7. Audiencias", 1)
    add_heading(doc, "Clientes principales", 2)
    for item in [
        "Terminales maritimos y terrestres.",
        "Traders.",
        "Navieras.",
        "Refinerias.",
        "Operadores logisticos de combustibles.",
        "Distribuidores.",
        "Aseguradoras y liquidadores.",
        "Clientes industriales con operaciones de combustibles.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Necesidades que resolvemos", 2)
    for item in [
        "Entender por que existe una diferencia de cantidad.",
        "Documentar eventos durante la operacion.",
        "Reducir demoras evitables.",
        "Respaldar reclamos o defensas tecnicas.",
        "Controlar calidad y trazabilidad documental.",
        "Tener ojos tecnicos independientes en terreno.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Identidad visual", 1)
    add_heading(doc, "Logo e isotipo", 2)
    add_para(doc, "El logo principal combina isotipo y wordmark ACI Loss Control Experts. Debe utilizarse en versiones de alto contraste, preferentemente sobre fondos negros, blancos o fotograficos con capa oscura.")
    add_para(doc, "El isotipo circular de ACI puede usarse como favicon, sello, marca de agua o elemento de apoyo. No reemplaza el logo completo en piezas comerciales principales salvo espacios reducidos.")
    add_heading(doc, "Paleta cromatica", 2)
    add_matrix(
        doc,
        ["Uso", "Color", "Hex"],
        [
            ("Negro operacional", "Fondo principal", "#0B0F12"),
            ("Blanco tecnico", "Texto y contraste", "#FFFFFF"),
            ("Gris acero", "Apoyo tecnico", "#5D788A"),
            ("Ambar operacional", "Acentos y llamados", "#C88C3A"),
            ("Verde industrial", "Bloques secundarios", "#546E57"),
            ("Azul maritimo", "Detalles y servicios", "#1F5867"),
        ],
        widths=(2.1, 2.6, 1.8),
    )
    add_heading(doc, "Estilo fotografico", 2)
    for item in [
        "Buques petroleros.",
        "Terminales maritimos.",
        "Plantas y tanques.",
        "Operaciones de carga y descarga.",
        "Instrumentacion, medicion, muestras y documentos.",
        "Escenas amplias, sobrias y profesionales.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9. Arquitectura de marca", 1)
    add_matrix(
        doc,
        ["Linea", "Descripcion"],
        [
            ("ACI Loss Control", "Presencia tecnica, control de eventos, seguimiento de operaciones y alertas tempranas."),
            ("ACI Expediting", "Seguimiento activo de tiempos, coordinaciones, hitos operacionales y demoras."),
            ("ACI Reconciliation", "Reconciliacion de cantidades, documentos, mediciones y diferencias."),
            ("ACI Claims Support", "Soporte tecnico-documental para reclamos, defensas, controversias y cierres comerciales."),
            ("ACI Reporting", "Reportes preliminares, bitacoras, dossier documental e informes finales."),
        ],
        widths=(2.0, 4.5),
    )
    add_heading(doc, "Regla de crecimiento", 2)
    for item in [
        "Tener procedimiento documentado.",
        "Tener personal competente o proveedor homologado.",
        "Tener formato de reporte y control de calidad interno.",
    ]:
        add_number(doc, item)

    add_heading(doc, "10. Aplicaciones iniciales", 1)
    add_kv_table(
        doc,
        [
            ("Web", "Quienes somos, problema que resolvemos, servicios, cobertura, manuales internos y contacto."),
            ("Presentacion comercial", "Portada, problema del cliente, propuesta ACI, servicios, cobertura, metodologia, reportabilidad y contacto."),
            ("Informes", "Resumen ejecutivo, datos de operacion, cronologia, hallazgos, diferencias, evidencias, conclusiones tecnicas y anexos."),
        ],
    )
    add_callout(
        doc,
        "Principio rector",
        "ACI debe verse, hablar y operar como una empresa que llega a terreno, entiende la operacion, documenta con precision y entrega al cliente informacion defendible.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
