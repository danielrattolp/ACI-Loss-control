from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "identidad"
LOGO = ROOT / "assets" / "aci-logo-black.jpeg"

INK = "0B0F12"
WHITE = "FFFFFF"
AMBER = "C88C3A"
SEA = "1F5867"
STEEL = "5D788A"
PAPER = "F6F7F4"
LINE = "D7DDE1"


def set_font(run, size=10.5, color="25313A", bold=False, italic=False, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    b = tbl_pr.first_child_found_in("w:tblBorders")
    if b is None:
        b = OxmlElement("w:tblBorders")
        tbl_pr.append(b)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = b.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}")
            b.append(e)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), size)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)


def base_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("ACI Loss Control Experts")
    set_font(run, size=8, color="707A82")
    return doc


def cover(doc, title, subtitle):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    borders(table, color=INK, size="0")
    cell = table.cell(0, 0)
    shade(cell, INK)
    margins(cell, top=240, start=260, bottom=240, end=260)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(3.35))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(20)
    r = p2.add_run(title)
    set_font(r, size=23, color=WHITE, bold=True)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    set_font(r3, size=11.5, color="DCE3E7")
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Control tecnico donde cada barril importa")
    set_font(r4, size=10, color=AMBER, bold=True)
    doc.add_page_break()


def h(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 12.5, color=SEA if level == 1 else STEEL, bold=True)


def p(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    set_font(run, bold=bold, italic=italic)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    set_font(run)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    borders(t)
    for i, header in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Inches(widths[i])
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(c, INK)
        margins(c)
        r = c.paragraphs[0].add_run(header)
        set_font(r, size=9, color=WHITE, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]
            c.width = Inches(widths[i])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            margins(c)
            if i == 0:
                shade(c, PAPER)
            r = c.paragraphs[0].add_run(value)
            set_font(r, size=9, bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def proposal_template():
    doc = base_doc()
    cover(doc, "Plantilla de Propuesta Comercial", "Documento base para presentar servicios ACI a clientes corporativos")
    h(doc, "1. Antecedentes del cliente")
    table(doc, ["Campo", "Detalle"], [
        ("Cliente", "[Nombre de la empresa]"),
        ("Contacto", "[Nombre, cargo y datos]"),
        ("Operacion", "[Carga / descarga / transferencia / almacenamiento]"),
        ("Ubicacion", "[Pais, ciudad, terminal, planta o nave]"),
        ("Producto", "[Combustible, crudo o derivado]"),
    ], [1.8, 4.7])
    h(doc, "2. Necesidad identificada")
    p(doc, "[Describir la exposicion operacional: diferencias de cantidad, control de calidad, tiempos, documentos, reclamos o necesidad de presencia tecnica independiente.]")
    h(doc, "3. Propuesta ACI")
    p(doc, "ACI propone ejecutar servicios de loss control petrolero y expediting con presencia tecnica, control de eventos, reconciliacion de cantidades, control documental y reportabilidad trazable.")
    h(doc, "4. Alcance del servicio")
    table(doc, ["Actividad", "Resultado esperado"], [
        ("Planificacion", "Revision de requerimientos, ubicacion, producto, tiempos y criterios de reporte."),
        ("Presencia tecnica", "Asistencia durante hitos criticos de la operacion."),
        ("Control de cantidad y calidad", "Supervision de mediciones, muestras, temperatura, aforo y certificados."),
        ("Reconciliacion", "Comparacion de datos entre nave, terminal, planta u operador."),
        ("Reportabilidad", "Reporte preliminar, bitacora, evidencias e informe final."),
    ], [2.0, 4.5])
    h(doc, "5. Entregables")
    for item in ["Reporte preliminar.", "Bitacora de eventos.", "Registro fotografico.", "Resumen de diferencias.", "Informe final con anexos documentales."]:
        bullet(doc, item)
    h(doc, "6. Condiciones comerciales")
    table(doc, ["Concepto", "Detalle"], [
        ("Honorarios", "[Valor / jornada / operacion / proyecto]"),
        ("Gastos reembolsables", "[Traslados, alojamiento, permisos u otros]"),
        ("Plazo de reporte", "[Definir SLA]"),
        ("Validez de propuesta", "[Dias]"),
    ], [2.0, 4.5])
    out = OUT_DIR / "ACI-Plantilla-Propuesta-Comercial.docx"
    doc.save(out)
    return out


def service_sheet():
    doc = base_doc()
    cover(doc, "Ficha Tecnica de Servicio", "Loss Control y Expediting")
    h(doc, "Servicio")
    p(doc, "Loss control petrolero y expediting para operaciones de carga, descarga, transferencia y almacenamiento de mercancias petroleras.")
    h(doc, "Objetivo")
    p(doc, "Proteger los intereses del cliente mediante presencia tecnica independiente, control de eventos, reconciliacion de cantidades, supervision documental y reportabilidad trazable.")
    h(doc, "Alcance tipico")
    for item in ["Revision de instrucciones operacionales.", "Coordinacion previa con partes involucradas.", "Presencia en hitos criticos.", "Registro cronologico de eventos.", "Supervision de mediciones, muestras y documentos.", "Alertas tempranas ante desviaciones.", "Informe final con evidencias."]:
        bullet(doc, item)
    h(doc, "Entregables")
    table(doc, ["Entregable", "Contenido"], [
        ("Reporte preliminar", "Resumen temprano de la operacion, eventos y diferencias relevantes."),
        ("Bitacora", "Linea de tiempo de hitos y comunicaciones."),
        ("Dossier", "Documentos, fotografias, certificados y respaldos."),
        ("Informe final", "Analisis tecnico, hallazgos, conclusiones y anexos."),
    ], [1.9, 4.6])
    h(doc, "Indicadores recomendados")
    table(doc, ["Indicador", "Meta referencial"], [
        ("Reporte preliminar dentro de plazo", "95%"),
        ("Informes sin reproceso documental", "98%"),
        ("Evidencias completas por operacion", "100%"),
        ("Acciones correctivas cerradas", "90%"),
    ], [3.2, 3.3])
    out = OUT_DIR / "ACI-Ficha-Tecnica-Loss-Control.docx"
    doc.save(out)
    return out


def logo_usage():
    doc = base_doc()
    cover(doc, "Guia de Uso de Logo y Marca", "Reglas practicas para aplicar la identidad ACI")
    h(doc, "Uso principal")
    p(doc, "El logo completo ACI Loss Control Experts debe utilizarse en portadas, propuestas, informes, brochure, sitio web y firmas comerciales.")
    h(doc, "Uso del isotipo")
    p(doc, "El isotipo debe utilizarse en favicon, sellos, marca de agua, iconos internos y aplicaciones donde el logo completo pierde legibilidad.")
    h(doc, "Versiones permitidas")
    table(doc, ["Version", "Uso recomendado"], [
        ("Blanco sobre negro", "Fondos oscuros, portadas, hero web y documentos premium."),
        ("Negro sobre blanco", "Documentos impresos, anexos, hojas internas y aplicaciones simples."),
        ("Isotipo", "Favicon, sello, icono, marca de agua y espacios reducidos."),
    ], [2.2, 4.3])
    h(doc, "Reglas de aplicacion")
    for item in ["Mantener contraste alto.", "No deformar el logo.", "No cambiar proporciones.", "No agregar sombras decorativas innecesarias.", "No usar sobre fondos con ruido visual.", "No combinar con logos de terceros sin autorizacion.", "Mantener margen de seguridad alrededor del logo."]:
        bullet(doc, item)
    h(doc, "Paleta")
    table(doc, ["Color", "Hex", "Uso"], [
        ("Negro operacional", "#0B0F12", "Fondos principales."),
        ("Blanco tecnico", "#FFFFFF", "Contraste y texto."),
        ("Ambar operacional", "#C88C3A", "Acentos y llamados."),
        ("Azul maritimo", "#1F5867", "Detalles tecnicos."),
        ("Gris acero", "#5D788A", "Apoyo visual."),
    ], [2.0, 1.5, 3.0])
    out = OUT_DIR / "ACI-Guia-Uso-Logo-Marca.docx"
    doc.save(out)
    return out


def client_onboarding():
    doc = base_doc()
    cover(doc, "Onboarding Comercial de Cliente", "Guia para levantar requerimientos y activar servicios ACI")
    h(doc, "Objetivo")
    p(doc, "Estandarizar el levantamiento inicial de clientes y operaciones para asegurar que ACI reciba la informacion necesaria antes de ejecutar servicios de loss control.")
    h(doc, "Informacion minima del cliente")
    table(doc, ["Dato", "Detalle requerido"], [
        ("Empresa", "[Razon social / nombre comercial]"),
        ("Contacto operacional", "[Nombre, cargo, telefono, email]"),
        ("Contacto administrativo", "[Nombre, cargo, telefono, email]"),
        ("Pais y ubicacion", "[Chile / Ecuador / Colombia - punto operacional]"),
        ("Facturacion", "[Datos administrativos]"),
    ], [2.0, 4.5])
    h(doc, "Informacion minima de la operacion")
    for item in ["Producto.", "Volumen estimado.", "Fecha y ventana operacional.", "Terminal, nave, planta o instalacion.", "Partes involucradas.", "Documentos disponibles.", "Criterios de reporte.", "Requisitos de acceso y seguridad operacional."]:
        bullet(doc, item)
    h(doc, "Flujo de activacion")
    table(doc, ["Paso", "Accion"], [
        ("1", "Recepcion del requerimiento."),
        ("2", "Revision de alcance y riesgos."),
        ("3", "Confirmacion de disponibilidad."),
        ("4", "Emision de propuesta u orden de servicio."),
        ("5", "Asignacion de inspector/coordinador."),
        ("6", "Ejecucion y reportabilidad."),
    ], [0.8, 5.7])
    h(doc, "Cierre inicial")
    p(doc, "Una vez aceptado el servicio, ACI debe registrar la orden de servicio, confirmar responsables, criterios de comunicacion, formato de reportes y plazos de entrega.")
    out = OUT_DIR / "ACI-Onboarding-Comercial-Cliente.docx"
    doc.save(out)
    return out


def comms_checklist():
    doc = base_doc()
    cover(doc, "Checklist de Comunicacion Corporativa", "Criterios para mantener consistencia en piezas ACI")
    h(doc, "Antes de publicar o enviar")
    for item in ["El mensaje explica que ACI hace loss control petrolero.", "El texto evita promesas absolutas.", "No se mencionan clientes especificos sin autorizacion.", "La pieza usa el logo correcto.", "La paleta respeta negro, blanco, ambar, azul maritimo y gris acero.", "La fotografia muestra operacion petrolera o maritima realista.", "El llamado a la accion es claro.", "El documento tiene datos de contacto actualizados."]:
        bullet(doc, item)
    h(doc, "Checklist por tipo de pieza")
    table(doc, ["Pieza", "Debe incluir"], [
        ("Web", "Propuesta de valor, servicios, cobertura y contacto."),
        ("Brochure", "Problema, solucion, servicios, metodologia y contacto."),
        ("Propuesta", "Alcance, entregables, plazos, condiciones y validez."),
        ("Informe", "Resumen ejecutivo, cronologia, hallazgos, evidencias y conclusiones."),
        ("LinkedIn", "Descripcion clara, especialidades y cobertura regional."),
    ], [1.7, 4.8])
    h(doc, "Frases aprobadas")
    for item in ["Control tecnico donde cada barril importa.", "Evidencia clara para operaciones criticas.", "Loss control petrolero con presencia regional.", "Trazabilidad, control y respuesta en terreno.", "Ojos tecnicos independientes para mercancias petroleras."]:
        bullet(doc, item)
    out = OUT_DIR / "ACI-Checklist-Comunicacion-Corporativa.docx"
    doc.save(out)
    return out


def main():
    for path in [proposal_template(), service_sheet(), logo_usage(), client_onboarding(), comms_checklist()]:
        print(path)


if __name__ == "__main__":
    main()
