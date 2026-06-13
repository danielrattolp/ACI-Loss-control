from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PATH = ROOT / "docs" / "manuales-calidad" / "ACI-SGC-00-Indice-Maestro.docx"


def font(run, size=10, color="25313A", bold=False):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Arial")
    rpr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def borders(table, color="D7DDE1"):
    tbl_pr = table._tbl.tblPr
    node = tbl_pr.first_child_found_in("w:tblBorders")
    if node is None:
        node = OxmlElement("w:tblBorders")
        tbl_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = OxmlElement(f"w:{edge}")
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "6")
        item.set(qn("w:color"), color)
        node.append(item)


doc = Document(PATH)
heading = doc.add_paragraph(style="Heading 1")
heading.add_run("Documentacion operacional vinculada")

rows = [
    ("ACI-OP-000", "Indice Maestro Operacional"),
    ("ACI-MO-001", "Manual Operacional de Loss Control y Expediting"),
    ("ACI-PT-001", "Control de Cantidad y Reconciliacion"),
    ("ACI-PT-002", "Muestreo, Calidad y Cadena de Custodia"),
    ("ACI-PT-003", "Expediting y Control de Tiempos"),
    ("ACI-PT-004", "Evidencia, Fotografias y Reportabilidad"),
    ("ACI-PT-005", "Investigacion de Diferencias y Soporte a Reclamos"),
    ("ACI-HSE-001", "Seguridad Operacional y Respuesta a Emergencias"),
    ("ACI-PC-001", "Capacitacion y Habilitacion Tecnica"),
    ("ACI-PQ-001", "Paquete de Formatos Operacionales"),
    ("ACI-PQ-002", "Control Documental y Competencia"),
]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True
borders(table)
for i, value in enumerate(("Codigo", "Documento", "Estado")):
    cell = table.rows[0].cells[i]
    shade(cell, "0B0F12")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    font(cell.paragraphs[0].add_run(value), size=9, color="FFFFFF", bold=True)
for index, (code, title) in enumerate(rows):
    cells = table.add_row().cells
    for cell in cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if index % 2:
        for cell in cells:
            shade(cell, "F6F7F4")
    font(cells[0].paragraphs[0].add_run(code), size=9, bold=True)
    font(cells[1].paragraphs[0].add_run(title), size=9)
    font(cells[2].paragraphs[0].add_run("Para aprobacion"), size=9)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
font(p.add_run("Condicion de vigencia: "), bold=True, color="1F5867")
font(p.add_run("los documentos operacionales requieren responsables nominales, aprobacion, capacitacion y prueba piloto antes de declararse vigentes."))

doc.save(PATH)
print(PATH)
