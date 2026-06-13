import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
LOGO = ROOT / "assets" / "aci-logo-black.jpeg"

INK = "0B0F12"
WHITE = "FFFFFF"
AMBER = "C88C3A"
SEA = "1F5867"
STEEL = "5D788A"
PAPER = "F6F7F4"
LINE = "D7DDE1"


NAME_MAP = {
    "estrategia-empresa.md": "ACI-Estrategia-Empresarial.docx",
    "perfil-corporativo.md": "ACI-Perfil-Corporativo.docx",
    "00-indice-sgc.md": "ACI-SGC-00-Indice-Maestro.docx",
    "01-manual-calidad.md": "ACI-SGC-01-Manual-de-Calidad.docx",
    "02-politica-calidad.md": "ACI-SGC-02-Politica-de-Calidad.docx",
    "03-objetivos-calidad.md": "ACI-SGC-03-Objetivos-de-Calidad.docx",
    "04-mapa-procesos.md": "ACI-SGC-04-Mapa-de-Procesos.docx",
    "05-matriz-riesgos-oportunidades.md": "ACI-SGC-05-Matriz-Riesgos-Oportunidades.docx",
    "06-procedimiento-control-documentos.md": "ACI-SGC-06-Procedimiento-Control-Documentos.docx",
    "07-procedimiento-operaciones-loss-control.md": "ACI-SGC-07-Procedimiento-Operaciones-Loss-Control.docx",
    "08-procedimiento-competencia-personal.md": "ACI-SGC-08-Procedimiento-Competencia-Personal.docx",
    "09-procedimiento-no-conformidades-acciones-correctivas.md": "ACI-SGC-09-Procedimiento-No-Conformidades.docx",
    "10-procedimiento-auditoria-interna.md": "ACI-SGC-10-Procedimiento-Auditoria-Interna.docx",
    "11-procedimiento-evaluacion-proveedores.md": "ACI-SGC-11-Procedimiento-Evaluacion-Proveedores.docx",
    "12-registros-formatos.md": "ACI-SGC-12-Registros-y-Formatos.docx",
}


def set_font(run, size=10.5, color="25313A", bold=False, italic=False, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color=LINE):
    tbl_pr = table._tbl.tblPr
    node = tbl_pr.first_child_found_in("w:tblBorders")
    if node is None:
        node = OxmlElement("w:tblBorders")
        tbl_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = node.find(qn(f"w:{edge}"))
        if item is None:
            item = OxmlElement(f"w:{edge}")
            node.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "6")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), color)


def base_doc(title, category):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.76)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for style_name, size, color, before, after in (
        ("Heading 1", 16, SEA, 14, 7),
        ("Heading 2", 13, SEA, 11, 5),
        ("Heading 3", 11.5, STEEL, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    set_font(header.add_run(f"ACI LATAM | {category}"), size=8, color=STEEL, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Documento de trabajo | Version inicial"), size=8, color=STEEL)

    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover.autofit = False
    cover.columns[0].width = Inches(6.5)
    borders(cover, INK)
    cell = cover.cell(0, 0)
    shade(cell, INK)
    margins(cell, top=260, start=240, bottom=260, end=240)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(2.8))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(18)
    set_font(p2.add_run(title.upper()), size=22, color=WHITE, bold=True)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p3.add_run(category), size=11, color="DCE3E7", bold=True)
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(12)
    set_font(p4.add_run("ACI LATAM | acilatam.cl"), size=9.5, color=AMBER, bold=True)
    doc.add_paragraph()
    meta = doc.add_table(rows=2, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    borders(meta)
    values = [["Version", "1.0", "Estado", "Documento de trabajo"], ["Fecha", "12-06-2026", "Responsable", "Por asignar"]]
    for i, row in enumerate(values):
        for j, value in enumerate(row):
            c = meta.cell(i, j)
            c.width = Inches(1.625)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            margins(c)
            if j % 2 == 0:
                shade(c, PAPER)
            set_font(c.paragraphs[0].add_run(value), size=9, bold=(j % 2 == 0))
    doc.add_page_break()
    return doc


def clean_inline(text):
    return text.replace("`", "").replace("**", "").replace("__", "").strip()


def add_rich_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    pieces = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            set_font(p.add_run(piece[2:-2]), bold=True)
        elif piece.startswith("`") and piece.endswith("`"):
            set_font(p.add_run(piece[1:-1]), size=9.5, name="Consolas", color=SEA)
        else:
            set_font(p.add_run(piece))
    return p


def add_md_table(doc, rows):
    if len(rows) < 2:
        return
    headers = [clean_inline(x) for x in rows[0]]
    body = rows[2:] if all(set(x.strip()) <= {"-", ":"} for x in rows[1]) else rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, INK)
        margins(cell)
        set_font(cell.paragraphs[0].add_run(header), size=8.8, color=WHITE, bold=True)
    for r_index, row in enumerate(body):
        cells = table.add_row().cells
        for idx in range(len(headers)):
            cell = cells[idx]
            margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_index % 2 == 1:
                shade(cell, PAPER)
            value = clean_inline(row[idx]) if idx < len(row) else ""
            set_font(cell.paragraphs[0].add_run(value), size=8.8, bold=(idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table_line(line):
    return [part.strip() for part in line.strip().strip("|").split("|")]


def convert(source, target):
    lines = source.read_text(encoding="utf-8").splitlines()
    first_heading = next((clean_inline(line.lstrip("# ")) for line in lines if line.startswith("# ")), source.stem.replace("-", " ").title())
    category = "Sistema de Gestion y Calidad" if "manuales-calidad" in source.parts else "Documentacion Comercial"
    doc = base_doc(first_heading, category)
    table_rows = []
    code_lines = []
    in_code = False

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_md_table(doc, table_rows)
            table_rows = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            borders(table, STEEL)
            cell = table.cell(0, 0)
            shade(cell, PAPER)
            margins(cell, top=130, start=160, bottom=130, end=160)
            text = "\n".join(code_lines)
            set_font(cell.paragraphs[0].add_run(text), size=8.5, name="Consolas", color=STEEL)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            code_lines = []

    skipped_title = False
    for line in lines:
        if line.strip().startswith("```"):
            flush_table()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.strip().startswith("|"):
            table_rows.append(parse_table_line(line))
            continue
        flush_table()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# ") and not skipped_title:
            skipped_title = True
            continue
        if stripped.startswith("### "):
            doc.add_heading(clean_inline(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(clean_inline(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(clean_inline(stripped[2:]), level=1)
        elif re.match(r"^\d+\.\s+", stripped):
            add_rich_paragraph(doc, re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            add_rich_paragraph(doc, stripped[2:], style="List Bullet")
        else:
            add_rich_paragraph(doc, stripped)
    flush_table()
    flush_code()
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return target


def main():
    sources = list((DOCS / "comercial").glob("*.md")) + list((DOCS / "manuales-calidad").glob("*.md"))
    outputs = []
    for source in sources:
        filename = NAME_MAP[source.name]
        outputs.append(convert(source, source.parent / filename))
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
