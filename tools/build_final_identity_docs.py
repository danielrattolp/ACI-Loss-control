from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "identidad"
LOGO = ROOT / "assets" / "aci-logo-black.jpeg"

INK = "0B0F12"
WHITE = "FFFFFF"
AMBER = "C88C3A"
SEA = "1F5867"
STEEL = "5D788A"
PAPER = "F6F7F4"
LINE = "D7DDE1"


def set_font(run, size=10.5, color="25313A", bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    b = tbl_pr.first_child_found_in("w:tblBorders")
    if b is None:
        b = OxmlElement("w:tblBorders")
        tbl_pr.append(b)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = b.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}")
            b.append(e)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), size)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)


def base_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("ACI Loss Control Experts")
    set_font(run, size=8, color="707A82")
    return doc


def cover(doc, title, subtitle):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.5)
    borders(t, color=INK, size="0")
    c = t.cell(0, 0)
    shade(c, INK)
    margins(c, top=240, start=260, bottom=240, end=260)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(3.35))
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(20)
    r = p2.add_run(title)
    set_font(r, size=23, color=WHITE, bold=True)
    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    set_font(r3, size=11.5, color="DCE3E7")
    p4 = c.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Control tecnico donde cada barril importa")
    set_font(r4, size=10, color=AMBER, bold=True)
    doc.add_page_break()


def h(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 12.5, color=SEA if level == 1 else STEEL, bold=True)


def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, bold=bold, italic=italic)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    borders(t)
    for i, header in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Inches(widths[i])
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(c, INK)
        margins(c)
        r = c.paragraphs[0].add_run(header)
        set_font(r, size=9, color=WHITE, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]
            c.width = Inches(widths[i])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            margins(c)
            if i == 0:
                shade(c, PAPER)
            r = c.paragraphs[0].add_run(value)
            set_font(r, size=9, bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def stationery():
    doc = base_doc()
    cover(doc, "Papeleria Corporativa", "Membrete, nota, cotizacion y orden de servicio")
    h(doc, "Membrete corporativo")
    table(doc, ["Elemento", "Uso"], [
        ("Encabezado", "Logo ACI arriba a la izquierda o centrado en portada formal."),
        ("Pie de pagina", "ACI Loss Control Experts | Chile | Ecuador | Colombia | sitio web."),
        ("Color", "Negro operacional, blanco tecnico y acento ambar."),
        ("Tipografia", "Arial para documentos editables; Georgia solo en piezas editoriales."),
    ], [1.7, 4.8])
    h(doc, "Plantilla de carta")
    para(doc, "Fecha: [dd-mm-aaaa]")
    para(doc, "Destinatario: [Nombre / Cargo / Empresa]")
    para(doc, "Asunto: [Asunto de la comunicacion]")
    para(doc, "Estimado/a [Nombre]:")
    para(doc, "[Cuerpo de la comunicacion con tono tecnico, claro y directo.]")
    para(doc, "Atentamente,")
    para(doc, "[Nombre Apellido] | [Cargo] | ACI Loss Control Experts")
    h(doc, "Documentos donde aplica")
    for item in ["Cartas comerciales.", "Cotizaciones.", "Ordenes de servicio.", "Comunicaciones de cierre.", "Minutas ejecutivas."]:
        bullet(doc, item)
    out = OUT_DIR / "ACI-Papeleria-Corporativa.docx"
    doc.save(out)
    return out


def business_card():
    doc = base_doc()
    cover(doc, "Tarjeta de Presentacion", "Contenido y reglas de aplicacion")
    h(doc, "Frente recomendado")
    table(doc, ["Zona", "Contenido"], [
        ("Logo", "ACI Loss Control Experts."),
        ("Nombre", "[Nombre Apellido]."),
        ("Cargo", "[Cargo]."),
        ("Contacto", "Telefono, email corporativo y web."),
    ], [1.6, 4.9])
    h(doc, "Reverso recomendado")
    table(doc, ["Elemento", "Contenido"], [
        ("Claim", "Control tecnico donde cada barril importa."),
        ("Servicios", "Loss control | Expediting | Reconciliacion | Reporting."),
        ("Cobertura", "Chile | Ecuador | Colombia."),
    ], [1.6, 4.9])
    h(doc, "Criterios visuales")
    for item in ["Fondo negro para version premium.", "Logo blanco sobre negro.", "Acento ambar solo para claim o linea divisoria.", "No saturar con iconos.", "Mantener lectura clara en tamano pequeno."]:
        bullet(doc, item)
    out = OUT_DIR / "ACI-Tarjeta-Presentacion.docx"
    doc.save(out)
    return out


def presentation_script():
    doc = base_doc()
    cover(doc, "Guion Presentacion Comercial", "Estructura base para futura presentacion PowerPoint")
    h(doc, "Estructura recomendada")
    table(doc, ["Slide", "Titulo", "Mensaje"], [
        ("1", "ACI Loss Control Experts", "Control tecnico donde cada barril importa."),
        ("2", "El problema", "Diferencias de cantidad, calidad, tiempo o documentos pueden transformarse en perdida o controversia."),
        ("3", "Nuestra funcion", "ACI actua como los ojos tecnicos independientes del cliente en terreno."),
        ("4", "Servicios", "Loss control, expediting, reconciliacion, control de calidad, soporte a reclamos e informes."),
        ("5", "Metodologia", "Planificar, ejecutar, registrar, reconciliar, reportar y cerrar."),
        ("6", "Cobertura", "Chile, Ecuador y Colombia con capacidad nacional en cada pais."),
        ("7", "Reportabilidad", "Reporte preliminar, bitacora, evidencias, dossier e informe final."),
        ("8", "Por que ACI", "Independencia, trazabilidad, oportunidad, rigor tecnico y confidencialidad."),
        ("9", "Activacion", "Levantamiento de requerimiento, propuesta, asignacion y ejecucion."),
        ("10", "Contacto", "Solicitud de cobertura y datos comerciales."),
    ], [0.7, 2.0, 3.8])
    h(doc, "Notas de estilo")
    for item in ["Usar fotografia operacional maritima o petrolera.", "Titulos como afirmaciones, no etiquetas genericas.", "No usar clientes nombrados sin autorizacion.", "Mantener paleta ACI.", "Evitar exceso de texto por slide."]:
        bullet(doc, item)
    out = OUT_DIR / "ACI-Guion-Presentacion-Comercial.docx"
    doc.save(out)
    return out


def domain_email():
    doc = base_doc()
    cover(doc, "Dominio y Correos Corporativos", "Recomendacion de estructura digital de marca")
    h(doc, "Dominio recomendado")
    table(doc, ["Opcion", "Uso"], [
        ("acilatam.cl", "Dominio principal corporativo recomendado y vigente."),
        ("www.acilatam.cl", "Version www redireccionable al dominio principal."),
        ("correo.acilatam.cl", "Subdominio opcional para documentacion tecnica del servicio de correo."),
        ("portal.acilatam.cl", "Subdominio opcional futuro para portal operacional o clientes."),
    ], [2.1, 4.4])
    h(doc, "Correos recomendados")
    table(doc, ["Correo", "Uso"], [
        ("contacto@acilatam.cl", "Contacto general, web y recepcion de oportunidades."),
        ("operaciones@acilatam.cl", "Coordinacion operacional y activacion de servicios."),
        ("administracion@acilatam.cl", "Facturacion y soporte administrativo."),
        ("danielratto@acilatam.cl", "Correo personal corporativo de direccion."),
    ], [2.8, 3.7])
    h(doc, "Reglas")
    for item in ["Usar solo correos corporativos en propuestas e informes.", "Evitar correos personales para comunicaciones con clientes.", "Centralizar reportes en una casilla operativa.", "Configurar firma corporativa uniforme.", "Activar autenticacion de dos factores para cuentas criticas."]:
        bullet(doc, item)
    out = OUT_DIR / "ACI-Dominio-Correos-Corporativos.docx"
    doc.save(out)
    return out


def master_index():
    doc = base_doc()
    cover(doc, "Indice Maestro de Identidad", "Mapa completo de documentos corporativos ACI")
    h(doc, "Documentos estrategicos")
    table(doc, ["Documento", "Funcion"], [
        ("ACI-Brand-Book", "Define esencia, tono, valores, mensajes y sistema base de marca."),
        ("ACI-Arquitectura-de-Marca", "Ordena lineas de servicio y reglas de crecimiento."),
        ("ACI-Lineamientos-Visuales", "Define logo, paleta, fotografia, iconografia y composicion."),
        ("ACI-Mensajes-Comerciales", "Centraliza claim, pitch, objeciones y frases aprobadas."),
    ], [2.7, 3.8])
    h(doc, "Documentos comerciales")
    table(doc, ["Documento", "Funcion"], [
        ("ACI-Brochure-Corporativo", "Presentacion breve de empresa y servicios."),
        ("ACI-Plantilla-Propuesta-Comercial", "Base para propuestas a clientes."),
        ("ACI-Ficha-Tecnica-Loss-Control", "Ficha del servicio principal."),
        ("ACI-Perfil-LinkedIn-Empresa", "Textos para presencia digital profesional."),
    ], [2.7, 3.8])
    h(doc, "Documentos de aplicacion")
    table(doc, ["Documento", "Funcion"], [
        ("ACI-Firma-Correo-Corporativa", "Formato de firma y aviso de confidencialidad."),
        ("ACI-Plantilla-Informe-Operacional", "Estructura base de informe de servicio."),
        ("ACI-Onboarding-Comercial-Cliente", "Levantamiento de clientes y operaciones."),
        ("ACI-Checklist-Comunicacion-Corporativa", "Control antes de publicar/enviar piezas."),
        ("ACI-Papeleria-Corporativa", "Membrete y carta base."),
        ("ACI-Tarjeta-Presentacion", "Contenido y reglas de tarjeta."),
        ("ACI-Guion-Presentacion-Comercial", "Estructura para futura presentacion PPT."),
        ("ACI-Dominio-Correos-Corporativos", "Estructura recomendada de dominio y casillas."),
    ], [2.7, 3.8])
    h(doc, "Estado")
    para(doc, "Con este paquete, la identidad de ACI queda lista para uso inicial comercial, digital y documental. La siguiente etapa recomendada es avanzar a manuales de procedimiento operacional y sistema de gestion interna.")
    out = OUT_DIR / "ACI-Indice-Maestro-Identidad.docx"
    doc.save(out)
    return out


def main():
    for output in [stationery(), business_card(), presentation_script(), domain_email(), master_index()]:
        print(output)


if __name__ == "__main__":
    main()
