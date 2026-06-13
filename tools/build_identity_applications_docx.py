from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "identidad"
LOGO = ROOT / "assets" / "aci-logo-black.jpeg"
HERO = ROOT / "assets" / "oil-tanker-hero.png"

INK = "0B0F12"
WHITE = "FFFFFF"
AMBER = "C88C3A"
SEA = "1F5867"
STEEL = "5D788A"
PAPER = "F6F7F4"
LINE = "D7DDE1"


def set_font(run, size=10.5, color="25313A", bold=False, italic=False, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    b = tbl_pr.first_child_found_in("w:tblBorders")
    if b is None:
        b = OxmlElement("w:tblBorders")
        tbl_pr.append(b)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = b.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}")
            b.append(e)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), size)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)


def doc_base():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("ACI Loss Control Experts")
    set_font(r, size=8, color="707A82")
    return doc


def cover(doc, title, subtitle, with_hero=False):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    borders(table, color=INK, size="0")
    cell = table.cell(0, 0)
    shade(cell, INK)
    margins(cell, top=230, start=240, bottom=230, end=240)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(3.4))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(20)
    r2 = p2.add_run(title)
    set_font(r2, size=24, color=WHITE, bold=True)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    set_font(r3, size=11.5, color="DCE3E7")
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Control tecnico donde cada barril importa")
    set_font(r4, size=10, color=AMBER, bold=True)
    if with_hero and HERO.exists():
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(str(HERO), width=Inches(6.5))
    doc.add_page_break()


def h(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    size = 16 if level == 1 else 12.5
    color = SEA if level == 1 else STEEL
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=True)


def p(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    r = para.add_run(text)
    set_font(r, bold=bold, italic=italic)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.15
    r = para.add_run(text)
    set_font(r)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    borders(t)
    for i, header in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Inches(widths[i])
        shade(c, INK)
        margins(c)
        r = c.paragraphs[0].add_run(header)
        set_font(r, size=9, color=WHITE, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]
            c.width = Inches(widths[i])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            margins(c)
            if i == 0:
                shade(c, PAPER)
            r = c.paragraphs[0].add_run(value)
            set_font(r, size=9, bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def callout(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.5)
    borders(t, color="D9C29A", size="8")
    c = t.cell(0, 0)
    shade(c, "FFF7E8")
    margins(c, top=140, start=180, bottom=140, end=180)
    r = c.paragraphs[0].add_run(title)
    set_font(r, size=10.5, color=INK, bold=True)
    para = c.add_paragraph()
    r2 = para.add_run(body)
    set_font(r2, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def brochure():
    doc = doc_base()
    cover(doc, "Brochure Corporativo", "Loss control petrolero, expediting y control documental", with_hero=True)
    h(doc, "Quienes somos")
    p(doc, "ACI Loss Control Experts es una firma especializada en loss control petrolero para operaciones de carga, descarga, transferencia y almacenamiento de combustibles, crudos y derivados.")
    p(doc, "Actuamos como los ojos tecnicos del cliente en terminales, buques, plantas y patios, registrando eventos, verificando mediciones, consolidando evidencias y emitiendo reportes trazables.")
    callout(doc, "Propuesta de valor", "Presencia tecnica independiente para reducir perdidas, controversias, demoras y decisiones sin evidencia.")
    h(doc, "Servicios principales")
    table(
        doc,
        ["Servicio", "Alcance"],
        [
            ("Loss control y expediting", "Control de eventos, tiempos, alertas tempranas y seguimiento operacional."),
            ("Reconciliacion", "Comparacion de cantidades, mediciones y documentos entre nave, terminal o planta."),
            ("Cantidad y calidad", "Supervision de mediciones, muestreo, temperatura, aforo y certificados."),
            ("Soporte a reclamos", "Linea de tiempo, evidencias, documentos criticos y calculos de diferencia."),
            ("Reportabilidad", "Reporte preliminar, bitacora, dossier documental e informe final."),
        ],
        [2.0, 4.5],
    )
    h(doc, "Cobertura")
    for item in ["Chile: cobertura nacional para puertos, terminales, refinerias, plantas y rutas de abastecimiento.", "Ecuador: operaciones portuarias, terminales de almacenamiento y cadenas de suministro.", "Colombia: operaciones costeras e interiores con soporte de inspeccion y control documental."]:
        bullet(doc, item)
    h(doc, "Por que ACI")
    for item in ["Independencia tecnica.", "Trazabilidad documental.", "Alertas tempranas.", "Criterio operacional.", "Informes claros y defendibles."]:
        bullet(doc, item)
    h(doc, "Contacto")
    p(doc, "Sitio web: https://acilatam.cl")
    p(doc, "Email comercial: contacto@acilatam.cl")
    out = OUT_DIR / "ACI-Brochure-Corporativo.docx"
    doc.save(out)
    return out


def linkedin_profile():
    doc = doc_base()
    cover(doc, "Perfil LinkedIn y Descripcion Corporativa", "Textos listos para perfil de empresa y redes profesionales")
    h(doc, "Nombre de pagina")
    p(doc, "ACI Loss Control Experts")
    h(doc, "Tagline")
    p(doc, "Loss control petrolero, expediting y control documental para operaciones criticas.")
    h(doc, "Descripcion corta")
    p(doc, "ACI Loss Control Experts protege los intereses del cliente durante operaciones petroleras criticas mediante presencia tecnica en terreno, reconciliacion de cantidades, control de calidad, control documental y reportabilidad trazable.")
    h(doc, "Descripcion para LinkedIn")
    p(doc, "ACI Loss Control Experts es una empresa especializada en loss control petrolero para operaciones de carga, descarga, transferencia y almacenamiento de combustibles, crudos y derivados.")
    p(doc, "Con presencia en Chile, Ecuador y Colombia, actuamos como los ojos tecnicos del cliente en terminales, buques, plantas y patios. Nuestro foco es registrar eventos, verificar mediciones, consolidar evidencias y entregar reportes claros para reducir perdidas, controversias y decisiones sin respaldo.")
    h(doc, "Especialidades")
    for item in ["Loss control petrolero", "Expediting", "Reconciliacion nave/terminal/planta", "Control de cantidad y calidad", "Control documental", "Soporte a reclamos", "Reportabilidad tecnica"]:
        bullet(doc, item)
    h(doc, "Publicacion de lanzamiento")
    p(doc, "Nace ACI Loss Control Experts, una firma creada para entregar control tecnico independiente en operaciones petroleras criticas. Nuestro proposito es dar certeza donde cada diferencia importa: cantidad, calidad, tiempos, documentos y evidencias.")
    p(doc, "Operamos con cobertura regional en Chile, Ecuador y Colombia, apoyando a terminales, traders, navieras, refinerias, operadores logisticos, aseguradoras y clientes industriales.")
    out = OUT_DIR / "ACI-Perfil-LinkedIn-Empresa.docx"
    doc.save(out)
    return out


def email_signature():
    doc = doc_base()
    cover(doc, "Firma de Correo Corporativa", "Formato base para comunicaciones ACI")
    h(doc, "Firma recomendada")
    table(
        doc,
        ["Elemento", "Contenido"],
        [
            ("Nombre", "[Nombre Apellido]"),
            ("Cargo", "[Cargo]"),
            ("Empresa", "ACI Loss Control Experts"),
            ("Telefono", "+56 [numero]"),
            ("Email", "[correo]@acilatam.cl"),
            ("Web", "https://acilatam.cl"),
            ("Claim", "Control tecnico donde cada barril importa."),
        ],
        [1.45, 5.05],
    )
    h(doc, "Version texto")
    p(doc, "[Nombre Apellido]")
    p(doc, "[Cargo] | ACI Loss Control Experts")
    p(doc, "Loss control petrolero, expediting y control documental")
    p(doc, "Chile | Ecuador | Colombia")
    p(doc, "M: +56 [numero] | E: [correo]@acilatam.cl")
    p(doc, "W: https://acilatam.cl")
    p(doc, "Control tecnico donde cada barril importa.")
    h(doc, "Aviso de confidencialidad sugerido")
    p(doc, "La informacion contenida en este correo y sus anexos puede ser confidencial y estar destinada exclusivamente a su receptor. Si usted no es el destinatario, por favor notifique al remitente y elimine el mensaje.")
    out = OUT_DIR / "ACI-Firma-Correo-Corporativa.docx"
    doc.save(out)
    return out


def report_template():
    doc = doc_base()
    cover(doc, "Plantilla Base de Informe Operacional", "Formato inicial para reportes de loss control")
    h(doc, "1. Datos de operacion")
    table(
        doc,
        ["Campo", "Detalle"],
        [
            ("Cliente", "[Nombre cliente]"),
            ("Operacion", "[Carga / descarga / transferencia / almacenamiento]"),
            ("Producto", "[Producto]"),
            ("Ubicacion", "[Terminal / nave / planta]"),
            ("Fecha", "[dd-mm-aaaa]"),
            ("Inspector ACI", "[Nombre]"),
        ],
        [1.8, 4.7],
    )
    h(doc, "2. Resumen ejecutivo")
    p(doc, "[Resumen breve de la operacion, alcance del servicio, hitos principales, diferencias relevantes y estado de cierre.]")
    h(doc, "3. Cronologia de eventos")
    table(
        doc,
        ["Hora", "Evento", "Evidencia"],
        [
            ("[hh:mm]", "[Evento observado]", "[Documento / foto / comunicacion]"),
            ("[hh:mm]", "[Evento observado]", "[Documento / foto / comunicacion]"),
        ],
        [1.0, 3.6, 1.9],
    )
    h(doc, "4. Hallazgos y desviaciones")
    table(
        doc,
        ["Hallazgo", "Impacto", "Accion / comentario"],
        [
            ("[Hallazgo]", "[Cantidad / calidad / tiempo / documento]", "[Accion recomendada o comentario tecnico]"),
        ],
        [2.3, 1.7, 2.5],
    )
    h(doc, "5. Conclusiones tecnicas")
    p(doc, "[Conclusiones del servicio, diferencias detectadas, respaldo documental y recomendaciones de cierre.]")
    h(doc, "6. Anexos")
    for item in ["Registro fotografico.", "Documentos de medicion.", "Certificados o reportes de laboratorio.", "Comunicaciones relevantes.", "Calculos de reconciliacion."]:
        bullet(doc, item)
    out = OUT_DIR / "ACI-Plantilla-Informe-Operacional.docx"
    doc.save(out)
    return out


def main():
    outputs = [brochure(), linkedin_profile(), email_signature(), report_template()]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
