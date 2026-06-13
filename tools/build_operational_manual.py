from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "operaciones" / "ACI-MO-001-Manual-Operacional-Loss-Control-Expediting.docx"
LOGO = ROOT / "assets" / "aci-logo-black.jpeg"

INK = "0B0F12"
WHITE = "FFFFFF"
AMBER = "C88C3A"
SEA = "1F5867"
STEEL = "5D788A"
PAPER = "F6F7F4"
LINE = "D7DDE1"
LIGHT = "E8EEF0"


def set_font(run, size=10.5, color="25313A", bold=False, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def base_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in (
        ("Heading 1", 16, SEA, 14, 7),
        ("Heading 2", 13, SEA, 10, 5),
        ("Heading 3", 11.5, STEEL, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("ACI LATAM | ACI-MO-001 | Manual Operacional")
    set_font(run, size=8, color=STEEL, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Version 1.0 | Para aprobacion e implementacion piloto")
    set_font(run, size=8, color=STEEL)
    return doc


def paragraph(doc, text, bold=False, italic=False, color="25313A", after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_font(r, bold=bold, italic=italic, color=color)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    set_font(p.add_run(text), size=10.2)
    return p


def numbered(doc, title, body):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{title}: ")
    set_font(r, bold=True)
    set_font(p.add_run(body))


def add_table(doc, headers, rows, widths, header_fill=INK):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_margins(cell)
        shade(cell, header_fill)
        set_font(cell.paragraphs[0].add_run(header), size=8.8, color=WHITE, bold=True)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_margins(cell)
            if row_index % 2 == 1:
                shade(cell, PAPER)
            set_font(cell.paragraphs[0].add_run(str(value)), size=8.8, bold=(idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def callout(doc, label, text, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.72)
    table_borders(table, color=LINE)
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell_margins(cell, top=130, start=160, bottom=130, end=160)
    p = cell.paragraphs[0]
    set_font(p.add_run(f"{label}: "), size=10, color=SEA, bold=True)
    set_font(p.add_run(text), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def cover(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.72)
    table_borders(table, color=INK, size="0")
    cell = table.cell(0, 0)
    shade(cell, INK)
    cell_margins(cell, top=360, start=300, bottom=360, end=300)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(3.1))
    title = cell.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    set_font(title.add_run("MANUAL OPERACIONAL"), size=24, color=WHITE, bold=True)
    sub = cell.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(sub.add_run("Loss Control y Expediting de Mercancias Petroleras"), size=13, color="DCE3E7", bold=True)
    claim = cell.add_paragraph()
    claim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    claim.paragraph_format.space_before = Pt(18)
    set_font(claim.add_run("Control tecnico donde cada barril importa"), size=10.5, color=AMBER, bold=True)

    doc.add_paragraph()
    add_table(doc, ["Codigo", "Version", "Estado", "Fecha"], [["ACI-MO-001", "1.0", "Para aprobacion e implementacion piloto", "12-06-2026"]], [1.35, 0.9, 3.05, 1.42])
    add_table(doc, ["Elaborado por", "Revisado por", "Aprobado por"], [["Responsable de Operaciones", "Responsable de Calidad", "Gerencia General"]], [2.24, 2.24, 2.24])
    paragraph(doc, "Este documento no declara certificacion ISO 9001 ni sustituye requisitos legales, contractuales, terminales, maritimos, ambientales o de seguridad aplicables en cada jurisdiccion.", italic=True, color=STEEL)
    doc.add_page_break()


def build():
    doc = base_document()
    cover(doc)

    doc.add_heading("Control del documento", level=1)
    add_table(doc, ["Version", "Fecha", "Descripcion del cambio", "Responsable"], [["1.0", "12-06-2026", "Emision inicial para aprobacion e implementacion piloto.", "Operaciones / Calidad"]], [0.8, 1.1, 3.65, 1.17])
    callout(doc, "Condicion de uso", "Antes de emitirlo como vigente, Gerencia debe aprobar el documento, asignar responsables nominales, confirmar los plazos contractuales y capacitar al personal involucrado.")

    doc.add_heading("1. Objeto", level=1)
    paragraph(doc, "Establecer el metodo de ACI LATAM para recibir, planificar, ejecutar, supervisar, documentar, reportar y cerrar servicios de loss control y expediting de mercancias petroleras, asegurando independencia tecnica, trazabilidad, comunicacion oportuna y evidencia verificable.")

    doc.add_heading("2. Alcance", level=1)
    paragraph(doc, "Aplica a servicios realizados en terminales, buques, refinerias, plantas, depositos, patios, rutas y puntos de transferencia en Chile, Ecuador, Colombia, Peru, Caribe y Estados Unidos, sujetos a disponibilidad operativa, requisitos locales y alcance contractual.")
    bullet(doc, "Operaciones de carga, descarga, transferencia, almacenamiento y despacho.")
    bullet(doc, "Control de cantidad, calidad, tiempos, eventos y documentacion.")
    bullet(doc, "Reconciliacion nave/terminal/planta y soporte tecnico a diferencias o reclamos.")
    bullet(doc, "Reportes preliminares, bitacoras, evidencia fotografica, dossier e informe final.")
    callout(doc, "Exclusion", "ACI no presta servicios de vetting. La inspeccion de condicion o seguridad de una nave solo se ejecutara cuando forme parte expresa del alcance de loss control y sin presentarse como aprobacion de vetting.", fill="FFF4E5")

    doc.add_heading("3. Principios operacionales", level=1)
    add_table(doc, ["Principio", "Aplicacion"], [
        ["Independencia", "Registrar hechos observables y separar hechos, calculos, opiniones y datos de terceros."],
        ["Trazabilidad", "Todo hallazgo debe vincularse con fecha, hora, fuente, documento, fotografia, medicion o comunicacion."],
        ["Oportunidad", "Las desviaciones criticas se comunican durante la operacion, sin esperar el informe final."],
        ["Confidencialidad", "La informacion se utiliza solo para el servicio y se comparte con destinatarios autorizados."],
        ["Seguridad", "Ninguna instruccion comercial justifica ejecutar una actividad insegura o no autorizada."],
        ["Competencia", "Solo personal evaluado y habilitado puede ejecutar o revisar actividades tecnicas."],
    ], [1.35, 5.37])

    doc.add_heading("4. Funciones y responsabilidades", level=1)
    add_table(doc, ["Rol", "Responsabilidades minimas"], [
        ["Gerencia", "Aprobar politicas, recursos, responsables, excepciones y contratos de alto riesgo."],
        ["Coordinador de Operaciones", "Revisar requerimiento, definir alcance, asignar personal, mantener comunicaciones y controlar plazos."],
        ["Inspector / Loss Controller", "Ejecutar instrucciones, registrar eventos, verificar documentos, alertar desviaciones y proteger evidencia."],
        ["Revisor Tecnico", "Comprobar coherencia de calculos, cronologia, evidencia, conclusiones y cumplimiento del alcance."],
        ["Calidad", "Controlar documentos, competencias, no conformidades, indicadores y acciones de mejora."],
        ["Administracion", "Controlar orden de servicio, datos contractuales, facturacion y archivo comercial."],
    ], [1.55, 5.17])

    doc.add_heading("5. Flujo del servicio", level=1)
    numbered(doc, "Recepcion", "Registrar el requerimiento en ACI-F-001, identificar cliente, producto, volumen, lugar, fechas, partes, idioma, entregables, plazos y restricciones.")
    numbered(doc, "Revision", "Confirmar que el alcance corresponde a competencias de ACI, verificar independencia, riesgos, requisitos locales y disponibilidad de personal.")
    numbered(doc, "Planificacion", "Asignar personal habilitado, emitir instrucciones, preparar formatos, contactos, equipos, EPP, accesos y plan de comunicaciones.")
    numbered(doc, "Ejecucion", "Aplicar ACI-F-002, mantener ACI-F-003 y ACI-F-004, observar hitos, mediciones, muestreo, tiempos y documentos.")
    numbered(doc, "Alerta", "Comunicar de inmediato desviaciones criticas, documentar destinatario, hora, respuesta e instruccion recibida.")
    numbered(doc, "Reporte preliminar", "Emitir ACI-F-005 con hechos relevantes, diferencias provisionales y documentos pendientes.")
    numbered(doc, "Revision tecnica", "Verificar integridad, trazabilidad, calculos, lenguaje y correspondencia con instrucciones antes del envio final.")
    numbered(doc, "Cierre", "Emitir ACI-F-006, organizar dossier, confirmar recepcion, registrar comentarios y abrir no conformidad cuando corresponda.")

    doc.add_heading("6. Recepcion y revision del requerimiento", level=1)
    doc.add_heading("6.1 Informacion minima", level=2)
    for item in [
        "Identificacion del cliente, contacto autorizado y destinatarios de reportes.",
        "Producto, cantidad estimada, unidad de medida y especificacion aplicable.",
        "Buque, terminal, planta, tanque, camion, barcaza u otra unidad involucrada.",
        "ETA, ventana, fecha de inicio, duracion estimada y disponibilidad requerida.",
        "Alcance tecnico, hitos presenciales, frecuencia de reportes y plazo final.",
        "Normas, metodos, tolerancias, instrucciones y documentos contractuales aplicables.",
        "Requisitos de acceso, seguridad, confidencialidad, idioma y facturacion.",
    ]:
        bullet(doc, item)
    doc.add_heading("6.2 Criterios de aceptacion", level=2)
    paragraph(doc, "El Coordinador acepta el servicio solo cuando el alcance es claro, existe personal competente, se conocen los riesgos relevantes y no hay conflicto de interes no gestionado. Toda limitacion debe comunicarse y quedar registrada antes de la movilizacion.")

    doc.add_heading("7. Planificacion y asignacion", level=1)
    add_table(doc, ["Control", "Evidencia requerida"], [
        ["Competencia", "Matriz vigente, experiencia aplicable, induccion y habilitacion para el servicio."],
        ["Independencia", "Declaracion del asignado y revision de relaciones o intereses relevantes."],
        ["Seguridad", "EPP, induccion, permisos, condiciones de acceso y contactos de emergencia."],
        ["Equipos", "Disponibilidad, identificacion y calibracion/verificacion cuando corresponda."],
        ["Comunicacion", "Grupo de destinatarios, canales, horarios y criterio de escalamiento."],
        ["Documentos", "Instruccion operativa, formatos vigentes, nominacion y antecedentes recibidos."],
    ], [1.45, 5.27])

    doc.add_heading("8. Ejecucion en terreno", level=1)
    doc.add_heading("8.1 Inicio", level=2)
    for item in [
        "Registrar llegada, acceso, induccion, participantes, condiciones iniciales y documentos disponibles.",
        "Confirmar identidad de tanques, lineas, producto, secuencia, unidades y puntos de medicion.",
        "Registrar mediciones iniciales, sellos, agua libre, temperatura y observaciones cuando formen parte del alcance.",
        "Confirmar que cualquier dato no observado directamente se identifique como proporcionado por terceros.",
    ]:
        bullet(doc, item)
    doc.add_heading("8.2 Durante la transferencia", level=2)
    for item in [
        "Mantener bitacora cronologica con hora local y zona horaria definida.",
        "Registrar inicio, detenciones, reinicios, cambios de tanque, caudales, presiones y termino cuando sean accesibles.",
        "Observar muestreo, sellado, identificacion y cadena de custodia segun instrucciones aplicables.",
        "Documentar eventos, discrepancias, protestas, restricciones, comunicaciones y acciones adoptadas.",
        "No alterar equipos, valvulas, sellos o documentos sin autorizacion y competencia expresa.",
    ]:
        bullet(doc, item)
    doc.add_heading("8.3 Termino", level=2)
    for item in [
        "Registrar mediciones finales, remanentes, stripping, drenajes y condicion final cuando corresponda.",
        "Obtener o solicitar documentos finales, certificados, tickets y hojas de medicion.",
        "Comparar cifras preliminares y comunicar diferencias superiores al criterio definido por el cliente.",
        "Confirmar pendientes documentales y hora estimada de entrega de reportes.",
    ]:
        bullet(doc, item)

    doc.add_heading("9. Control de cantidad", level=1)
    paragraph(doc, "ACI verifica la consistencia de las mediciones y calculos conforme al metodo definido por el cliente, contrato, terminal, laboratorio o norma aplicable. El personal no debe inventar factores ni sustituir criterios normativos por practicas no documentadas.")
    add_table(doc, ["Punto de control", "Verificacion"], [
        ["Identificacion", "Tanque, compartimiento, producto, unidad, tabla y documento correcto."],
        ["Medicion", "Nivel, vacio/ullage, agua libre, temperatura, densidad y hora."],
        ["Instrumento", "Identificacion, condicion, rango y evidencia de calibracion/verificacion disponible."],
        ["Calculo", "Unidades, conversiones, correcciones, redondeos y factores aplicados."],
        ["Comparacion", "Diferencia absoluta, porcentual y respecto de tolerancia/instruccion del cliente."],
        ["Revision", "Segundo control independiente para calculos que sustenten diferencias o reclamos."],
    ], [1.35, 5.37])
    callout(doc, "Regla", "Las cifras deben distinguir claramente entre observadas por ACI, calculadas por ACI y proporcionadas por terceros.")

    doc.add_heading("10. Control de calidad y muestras", level=1)
    for item in [
        "Confirmar plan o instruccion de muestreo aplicable antes de iniciar.",
        "Registrar punto, metodo, fecha, hora, responsable, recipiente, sello y destino de cada muestra observada.",
        "Mantener identificacion unica y cadena de custodia cuando ACI reciba o traslade muestras.",
        "Comparar certificados y resultados con la especificacion suministrada, sin emitir conclusiones fuera de competencia.",
        "Informar contaminacion aparente, sellos alterados, recipientes inadecuados o perdida de trazabilidad.",
    ]:
        bullet(doc, item)

    doc.add_heading("11. Expediting y control de tiempos", level=1)
    paragraph(doc, "El expediting busca visibilidad y anticipacion. ACI registra hitos, identifica causas de demora, solicita aclaraciones y comunica impactos potenciales, sin sustituir la autoridad operacional de nave, terminal o cliente.")
    add_table(doc, ["Hito", "Registro minimo"], [
        ["Arribo / disponibilidad", "ETA, NOR u otro hito aplicable, restricciones y fuente de informacion."],
        ["Acceso / preparacion", "Permisos, inducciones, inspecciones, reunion y readiness."],
        ["Inicio", "Hora, condiciones, tanque/linea y observaciones."],
        ["Interrupcion", "Inicio, termino, duracion, causa reportada, parte responsable y evidencia."],
        ["Reinicio", "Hora, autorizacion y condicion relevante."],
        ["Termino / documentos", "Fin de transferencia, mediciones, documentacion y salida."],
    ], [1.55, 5.17])

    doc.add_heading("12. Alertas y escalamiento", level=1)
    add_table(doc, ["Nivel", "Ejemplos", "Accion"], [
        ["Critico", "Riesgo de seguridad, contaminacion, perdida significativa, instruccion contradictoria o integridad de evidencia comprometida.", "Detener la actividad propia si es insegura, informar inmediatamente al Coordinador y cliente autorizado, registrar hora y respuesta."],
        ["Alto", "Diferencia superior a tolerancia, medicion cuestionable, muestra sin trazabilidad, demora relevante o documento esencial faltante.", "Alertar durante la operacion y solicitar instruccion o aclaracion."],
        ["Medio", "Inconsistencia documental, retraso menor, dato pendiente o desviacion corregible.", "Registrar, comunicar en actualizacion programada y dar seguimiento."],
        ["Bajo", "Observacion sin impacto inmediato.", "Registrar para informe y mejora."],
    ], [0.75, 3.05, 2.92])

    doc.add_heading("13. Evidencia y control documental", level=1)
    paragraph(doc, "Cada servicio debe contar con una carpeta unica identificada por numero de orden. La estructura minima recomendada es: 01-Instrucciones, 02-Operacional, 03-Mediciones, 04-Calidad, 05-Fotografias, 06-Comunicaciones, 07-Reportes y 08-Cierre.")
    add_table(doc, ["Elemento", "Regla de trazabilidad"], [
        ["Nombre de archivo", "OS-AAAA-NNN_Tipo_FechaHora_Descripcion_Version.ext"],
        ["Fotografia", "Numero, fecha/hora, ubicacion, objeto, descripcion y autor en ACI-F-004."],
        ["Documento recibido", "Fuente, fecha/hora de recepcion y version si existe."],
        ["Correccion", "No eliminar el registro original; emitir version corregida y documentar el motivo."],
        ["Respaldo", "Copiar al repositorio autorizado tan pronto como la conectividad y seguridad lo permitan."],
        ["Acceso", "Restringido a personal asignado y destinatarios autorizados."],
    ], [1.4, 5.32])

    doc.add_heading("14. Reportabilidad", level=1)
    add_table(doc, ["Documento", "Contenido minimo", "Plazo"], [
        ["Actualizacion operacional", "Estado, hitos, eventos, diferencias preliminares y proximos pasos.", "Segun instruccion o evento critico."],
        ["ACI-F-005 Reporte preliminar", "Resumen, cifras provisionales, hechos, alertas, pendientes y evidencia disponible.", "Dentro del plazo contractual; si no existe, proponer 12 horas desde el termino."],
        ["ACI-F-006 Informe final", "Alcance, metodologia, cronologia, cantidades, calidad, tiempos, hallazgos, conclusiones y anexos.", "Dentro del plazo contractual; si no existe, proponer 3 dias habiles."],
    ], [1.55, 3.75, 1.42])
    paragraph(doc, "Los plazos indicados como propuesta deben confirmarse con cada cliente y registrarse en la orden de servicio.", italic=True, color=STEEL)

    doc.add_heading("15. Revision tecnica del informe", level=1)
    for item in [
        "El alcance y las limitaciones coinciden con la orden de servicio.",
        "Las cifras, unidades, formulas y redondeos fueron comprobados.",
        "La cronologia concuerda con bitacora, fotografias y comunicaciones.",
        "Cada hallazgo relevante posee evidencia identificable.",
        "Se distingue informacion observada, calculada y recibida de terceros.",
        "Las conclusiones son tecnicas, prudentes y no exceden la evidencia disponible.",
        "Se eliminaron datos innecesarios, comentarios informales y metadatos no autorizados.",
        "La version, aprobacion y destinatarios son correctos.",
    ]:
        bullet(doc, item)

    doc.add_heading("16. Diferencias, protestas y soporte a reclamos", level=1)
    paragraph(doc, "Ante una diferencia relevante, ACI preserva documentos, calcula de forma reproducible, identifica supuestos, construye una cronologia y evita atribuir responsabilidad legal. Cualquier carta de protesta u observacion debe registrarse y gestionarse conforme a instrucciones del cliente.")
    bullet(doc, "No firmar aceptacion de responsabilidad en nombre del cliente.")
    bullet(doc, "No modificar documentos de terceros; registrar observaciones por medio autorizado.")
    bullet(doc, "Mantener copia de documentos firmados, rechazados o recibidos con reserva.")
    bullet(doc, "Escalar posibles conflictos legales, regulatorios o reputacionales a Gerencia.")

    doc.add_heading("17. Seguridad, ambiente y autoridad para detener", level=1)
    paragraph(doc, "El personal ACI debe cumplir requisitos del sitio, utilizar EPP, respetar zonas restringidas y reportar condiciones inseguras. Puede suspender su propia actividad cuando exista riesgo no controlado, falta de autorizacion, equipo inadecuado o instruccion contraria a seguridad. ACI no dirige la operacion del terminal o buque.")

    doc.add_heading("18. Cierre y conservacion", level=1)
    for item in [
        "Confirmar entrega y recepcion de reportes por destinatarios autorizados.",
        "Verificar que el dossier contiene instrucciones, formatos, evidencia, comunicaciones e informes finales.",
        "Registrar documentos pendientes y responsable de seguimiento.",
        "Aplicar el plazo de retencion contractual o legal; en ausencia, ACI propone cinco anos sujeto a aprobacion.",
        "Registrar satisfaccion, reclamos, no conformidades y lecciones aprendidas.",
    ]:
        bullet(doc, item)

    doc.add_heading("19. Competencia y habilitacion", level=1)
    paragraph(doc, "Ninguna persona se considera habilitada solo por experiencia declarada. Debe existir evidencia de perfil, evaluacion, induccion, supervision inicial y autorizacion para los servicios asignados.")
    add_table(doc, ["Nivel", "Criterio", "Autorizacion"], [
        ["En formacion", "Induccion completada; experiencia o evaluacion aun insuficiente.", "Solo acompanado y sin firmar informes."],
        ["Habilitado", "Competencias evaluadas y desempeno satisfactorio en servicio supervisado.", "Ejecuta alcance definido."],
        ["Senior", "Experiencia demostrada, criterio de escalamiento y revision consistente.", "Lidera operaciones complejas y apoya formacion."],
        ["Revisor tecnico", "Dominio metodologico, independencia del trabajo revisado y autorizacion formal.", "Aprueba tecnicamente reportes."],
    ], [1.05, 3.4, 2.27])

    doc.add_heading("20. Indicadores de control", level=1)
    add_table(doc, ["Indicador", "Formula", "Meta inicial"], [
        ["Reportes preliminares oportunos", "Reportes dentro del plazo / reportes emitidos x 100", ">= 95%"],
        ["Informes aprobados sin reproceso mayor", "Informes sin devolucion mayor / informes revisados x 100", ">= 90%"],
        ["Evidencia completa", "Servicios con dossier completo / servicios cerrados x 100", "100%"],
        ["Personal habilitado", "Asignaciones con competencia vigente / asignaciones x 100", "100%"],
        ["Cierre de acciones", "Acciones cerradas en plazo / acciones vencidas en periodo x 100", ">= 90%"],
    ], [2.05, 3.25, 1.42])

    doc.add_heading("21. Registros asociados", level=1)
    add_table(doc, ["Codigo", "Registro", "Responsable primario"], [
        ["ACI-F-001", "Orden de servicio", "Coordinador de Operaciones"],
        ["ACI-F-002", "Checklist operacional", "Inspector / Loss Controller"],
        ["ACI-F-003", "Bitacora de eventos", "Inspector / Loss Controller"],
        ["ACI-F-004", "Registro fotografico", "Inspector / Loss Controller"],
        ["ACI-F-005", "Reporte preliminar", "Inspector y Coordinador"],
        ["ACI-F-006", "Informe final", "Inspector, Revisor Tecnico y Coordinador"],
        ["ACI-F-101", "Listado maestro de documentos", "Calidad"],
        ["ACI-F-201", "Perfil de cargo", "Gerencia / Personas"],
        ["ACI-F-202", "Matriz de competencia", "Operaciones / Calidad"],
        ["ACI-F-203", "Registro de capacitacion", "Personas / Calidad"],
    ], [1.05, 3.7, 1.97])

    doc.add_heading("22. Implementacion piloto", level=1)
    callout(doc, "Antes del primer servicio", "Aprobar este manual, completar los datos corporativos, nombrar responsables, capacitar al equipo, ejecutar una operacion simulada, corregir los formatos y emitir las versiones vigentes.", fill="FFF4E5")
    for item in [
        "Simular recepcion, planificacion, bitacora, alertas, reporte preliminar e informe final.",
        "Verificar que los archivos puedan completarse desde terreno y revisarse en oficina.",
        "Medir tiempo de elaboracion y detectar campos innecesarios o faltantes.",
        "Registrar hallazgos como acciones de mejora y aprobar la version 1.0 para uso real.",
    ]:
        bullet(doc, item)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
